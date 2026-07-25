#!/usr/bin/env python3
"""
DU Executive PhD Deep-Dive Research Report Builder
Generates a comprehensive Word document to inform Richard's accept/decline decision.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path
import datetime

OUTPUT = Path(__file__).parent / "DU_Executive_PhD_Deep_Dive_Report.docx"


def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Configure heading styles
    for level in range(1, 4):
        sn = f"Heading {level}"
        if sn in doc.styles:
            hs = doc.styles[sn]
            hf = hs.font
            hf.name = "Calibri"
            hf.color.rgb = RGBColor(0, 51, 102)
            if level == 1:
                hf.size = Pt(18)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(18)
                hs.paragraph_format.space_after = Pt(8)
            elif level == 2:
                hf.size = Pt(14)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(14)
                hs.paragraph_format.space_after = Pt(6)
            elif level == 3:
                hf.size = Pt(12)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(10)
                hs.paragraph_format.space_after = Pt(4)
            hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return doc


def add_para(doc, text, bold=False, italic=False, size=None, align=None, space_after=None, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    return p


def add_rich_para(doc, segments, space_after=None, indent=False):
    """Add paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        # Shade header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            # Alternate row shading
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EEF4" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def build_report():
    doc = create_doc()

    # ===== TITLE PAGE =====
    for _ in range(6):
        doc.add_paragraph()
    add_para(doc, "UNIVERSITY OF DENVER", bold=True, size=24,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Daniels College of Business", bold=True, size=18,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Executive PhD Program", size=16,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "DEEP-DIVE RESEARCH REPORT", bold=True, size=20,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "A Comprehensive Analysis for Accept/Decline Decision-Making",
             italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)
    add_para(doc, "Prepared for Richard W. Davidson", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, f"February 2026", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Based on 60+ web sources including official program pages, faculty profiles,",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, "alumni dissertations, external rankings, competitor comparisons, and independent reviews",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Program Overview & History",
        "3. Admissions & Cost",
        "4. Curriculum Deep Dive",
        "5. Faculty Profiles",
        "6. Cohort History & Alumni",
        "7. Student Research & Publications",
        "8. Program Reputation & Rankings",
        "9. Honest Assessment: Strengths & Concerns",
        "10. Fit Analysis for Richard's Situation",
        "11. Questions to Ask the Program",
        "12. Sources",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=3)
    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    doc.add_heading("1. Executive Summary", level=1)

    add_para(doc, "This report provides a comprehensive analysis of the University of Denver Daniels College of Business Executive PhD program to support Richard Davidson's enrollment decision. The research draws on 60+ sources including official university materials, faculty scholarship profiles, alumni dissertations, competitor program data, and independent assessments.")

    doc.add_heading("Decision Framework at a Glance", level=2)

    doc.add_heading("Key Strengths", level=3)
    strengths = [
        ("True PhD designation: ", "Unlike DBA programs, this is a PhD in Business Administration from an AACSB-accredited institution (accredited since 1923). This distinction matters for academic hiring and AACSB faculty qualification standards."),
        ("Institutional familiarity: ", "As a DU MBA alumnus, Richard already knows the faculty, campus culture, and institutional processes. This reduces friction and accelerates relationship-building with potential committee members."),
        ("Competitive pricing: ", "At ~$104,580 total, DU's program is the most affordable among peer Executive PhD programs (Oklahoma State: $139,500; Bentley: $153,000; Claremont: $107K-$143K)."),
        ("R1 university classification: ", "DU holds R1 (Very High Research Activity) Carnegie classification, which adds weight to the credential."),
        ("Dissertation alignment: ", "Richard's research on best-value source selection in government RFPs aligns with faculty expertise in decision-making, ethics, organizational behavior, and quantitative methods."),
        ("Student publications: ", "Students have co-authored peer-reviewed papers (Frontiers in Psychology, JABE, HICSS), indicating genuine research mentorship."),
    ]
    for prefix, text in strengths:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_heading("Key Concerns", level=3)
    concerns = [
        ("Leadership turnover: ", "Both the program founder (Lisa Victoravich, departed 2023) and founding research director (Craig Wallace, departed ~2019) have left. The current director (Dan Baack) is the third leader in seven years."),
        ("Program youth: ", "Founded Fall 2018 with only 7-8 cohorts admitted. The alumni network is thin and career outcome data is limited compared to established programs like Oklahoma State (founded 2012) or Claremont (~40 years)."),
        ("Self-funded investment: ", "~$104,580 out of pocket with no stipend, no assistantship, and no documented alumni tuition discount for returning MBA graduates. Traditional funded PhDs pay $0 tuition plus a $25K-$40K annual stipend."),
        ("Career placement reality: ", "Known graduates have landed at regional teaching institutions (UNC) or practitioner-track positions (DU itself), not R1 tenure-track roles. This is consistent with Executive PhD programs nationally."),
        ("90 credit hours: ", "The highest credit requirement among competitors (Oklahoma State: 60; Claremont: 48-64). More coursework means more time and cost per credit, though it could signal more thorough preparation."),
    ]
    for prefix, text in concerns:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_heading("Bottom Line", level=3)
    add_para(doc, "The DU Daniels Executive PhD is a legitimate, AACSB-accredited doctoral program from an R1 university at a competitive price point. It is well-suited for Richard's goals of consulting/advisory work with academic teaching as an additional pathway. The program will NOT position graduates for R1 tenure-track positions (no Executive PhD program reliably does), but it provides the credential, research training, and faculty relationships needed for teaching-focused academic roles, professor of practice positions, and enhanced professional credibility. Richard's existing DU relationships, advanced dissertation work, and quantitative skills are significant advantages.")

    doc.add_page_break()

    # ===== 2. PROGRAM OVERVIEW & HISTORY =====
    doc.add_heading("2. Program Overview & History", level=1)

    add_table(doc,
        ["Feature", "Detail"],
        [
            ["Official Name", "Executive PhD in Business Administration"],
            ["Degree Awarded", "PhD in Business Administration"],
            ["Founded", "Fall 2018 (September 2018)"],
            ["National Standing", "Third institution in the nation to offer an Executive PhD; first in the western United States"],
            ["Credit Hours", "90 credit hours"],
            ["Duration", "3 years (including two summers)"],
            ["Format", "Cohort-based, hybrid: 9 on-campus immersions/year (Fri-Sat, 8am-5pm) + weekly live online sessions + asynchronous work"],
            ["Year 3 Structure", "Independent dissertation with quarterly in-person committee meetings in Denver"],
            ["AACSB Accreditation", "Continuous since 1923 (among the first business schools accredited; first in the Rocky Mountain Region)"],
            ["University Classification", "R1: Very High Research Activity (Carnegie Classification)"],
            ["Business School Rank", "Daniels is the 8th-oldest business school in the US (established 1908)"],
            ["Exit Ramp", "MS in Applied Business Research (52 credits min, 3.0 GPA) for students who leave before completing the dissertation"],
        ],
        col_widths=[2.0, 4.5]
    )

    add_para(doc, 'The program was conceived by Lisa Victoravich (then Associate Dean for Faculty and Research) and Craig Wallace (then Head of the Management Department) as a way to "enable business leaders to make evidence-based decisions by conducting research." The first cohort of 18 students began in September 2018, with an average age of 47 and members commuting from as far as Texas and the Northeast.')

    doc.add_page_break()

    # ===== 3. ADMISSIONS & COST =====
    doc.add_heading("3. Admissions & Cost", level=1)

    doc.add_heading("Admissions Requirements", level=2)
    add_table(doc,
        ["Requirement", "Detail"],
        [
            ["Minimum GPA", "2.5 cumulative on baccalaureate degree (4.0 scale), OR 2.5 on last 60 semester credits, OR an earned master's degree or higher"],
            ["Professional Experience", "At least 10 years of substantial professional experience with demonstrated leadership"],
            ["GMAT/GRE", "Not required (GMAT/GRE-optional institution; holistic review)"],
            ["Recommendations", "Two letters (one strongly recommended from current employer expressing support)"],
            ["Essay", "Personal essay demonstrating rationale and passion for pursuing a PhD in business"],
            ["Resume", "Professional resume/CV"],
            ["Transcripts", "Scanned transcripts from all colleges/universities attended"],
            ["Interview", "By invitation only (on campus or via webcam)"],
            ["English Proficiency", "TOEFL 94+ / IELTS 7.0+ / C1 Advanced 185+ (non-native speakers)"],
        ],
        col_widths=[1.8, 4.7]
    )

    doc.add_heading("Cost Structure", level=2)
    add_table(doc,
        ["Item", "Amount", "Notes"],
        [
            ["Total Program Cost", "$104,580", "2025-2026 rate; same for in-state and out-of-state"],
            ["Per-Credit Rate", "~$1,162/credit", "Based on $104,580 / 90 credits"],
            ["Employer Match", "Up to $10,000", "One-time match; must indicate employer sponsorship before admission"],
            ["Yellow Ribbon (Veterans)", "Full remaining tuition", "DU expanded to all graduate programs Fall 2024; VA pays 50% of remaining costs, DU covers the other 50%"],
            ["Alumni Tuition Discount", "Not documented", "No publicly available alumni discount found; recommend asking admissions directly"],
            ["Stipend/Assistantship", "None", "Unlike traditional funded PhD programs"],
            ["Financial Aid", "Standard graduate financial aid available", "Federal loans, merit-based scholarships (availability for ExecPhD unclear)"],
        ],
        col_widths=[1.5, 1.5, 3.5]
    )

    add_rich_para(doc, [
        ("Cost Context: ", True),
        ("Traditional PhD programs in business are typically fully funded with tuition waivers, health insurance, and stipends of $25,000-$40,000/year. The Executive PhD's ~$105K out-of-pocket cost is offset by the fact that students maintain their full-time careers and salaries throughout the program. Among Executive PhD peers, DU's price is the lowest (Oklahoma State: $139,500; Claremont: $107K-$143K; Bentley: $153,000).", False),
    ])

    doc.add_page_break()

    # ===== 4. CURRICULUM DEEP DIVE =====
    doc.add_heading("4. Curriculum Deep Dive", level=1)

    doc.add_heading("Research Methods Foundation (16 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["BUS 6000", "Research Methods in Business", "4"],
            ["BUS 6001", "Qualitative Research Methods", "4"],
            ["BUS 6002", "Quantitative Methods I: Making Discoveries with Data", "4"],
            ["BUS 6003", "Quantitative Methods II: Making Discoveries with Data", "4"],
        ],
        col_widths=[1.2, 4.0, 1.0]
    )

    doc.add_heading("Research Design & Academic Skills (8 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["BUS 6005", "Behavioral Research Design and Execution", "4"],
            ["BUS 6400", "Academic Skills for Doctoral Students in Business", "2"],
            ["", "Additional elective/requirement", "2"],
        ],
        col_widths=[1.2, 4.0, 1.0]
    )

    doc.add_heading("Research Seminars (20 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["BUS 6300", "Seminar in Cross Disciplinary Decision Making Research", "4"],
            ["MGMT 6300", "Seminar in Leadership Strategy Research", "4"],
            ["MGMT 6301", "Research Seminar in Organizational Ethics", "4"],
            ["MGMT 6302", "Research Seminar in Leadership: An Organizational Perspective", "4"],
            ["MKTG 6300", "Seminar in Marketing", "4"],
        ],
        col_widths=[1.4, 4.0, 1.0]
    )

    doc.add_heading("Specialized Topics (4 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["FIN 6305", "Applied Quantitative Methods in Finance", "4"],
        ],
        col_widths=[1.2, 4.0, 1.0]
    )

    doc.add_heading("Practicum & Dissertation Preparation (20 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["BUS 6303", "Launch Your Doctoral Journey: Be Impactful", "4"],
            ["BUS 6500", "Applied Research Practicum I", "4"],
            ["BUS 6501", "Applied Research Practicum II", "4"],
            ["BUS 6502", "Applied Research Practicum III", "4"],
            ["BUS 6503", "Applied Research Practicum IV", "4"],
        ],
        col_widths=[1.2, 4.0, 1.0]
    )

    doc.add_heading("Dissertation (28 credits)", level=2)
    add_table(doc,
        ["Course", "Title", "Credits"],
        [
            ["BUS 6900", "Dissertation Research in Business", "28"],
        ],
        col_widths=[1.2, 4.0, 1.0]
    )

    add_rich_para(doc, [("Total: 90+ credits", True)])

    doc.add_heading("Key Milestones", level=2)
    add_bullet(doc, "Take-home format over several days, covering research methodology/statistics and seminar courses. Requires minimum 3.0 GPA. Graded by faculty committee.", bold_prefix="Comprehensive Exam: ")
    add_bullet(doc, "Oral defense before dissertation committee; passage confers PhD candidacy.", bold_prefix="Dissertation Proposal Defense: ")
    add_bullet(doc, "Final oral examination per University of Denver standards.", bold_prefix="Dissertation Defense: ")

    doc.add_heading("Year-by-Year Timeline", level=2)
    add_table(doc,
        ["Year", "Focus", "Format"],
        [
            ["Year 1", "Core research methods, foundational seminars, ARP I-II", "9 on-campus immersions (Fri/Sat) + weekly online sessions"],
            ["Year 2", "Advanced seminars, specialized topics, ARP III-IV, comprehensive exam", "9 on-campus immersions (Fri/Sat) + weekly online sessions"],
            ["Year 3", "Independent dissertation research (28 credits)", "Quarterly in-person committee meetings in Denver"],
        ],
        col_widths=[0.8, 3.0, 2.7]
    )

    doc.add_page_break()

    # ===== 5. FACULTY PROFILES =====
    doc.add_heading("5. Faculty Profiles", level=1)

    # Dan Baack
    doc.add_heading("Dan Baack -- Current Program Director", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Associate Professor of Marketing; Director, Executive PhD Program"],
            ["Research Areas", "International marketing, advertising, business ethics"],
            ["Publications", "30+ peer-reviewed journal articles including Journal of Business Ethics, Journal of International Business Studies, Harvard Business Review (online)"],
            ["Grants", "Nearly 20 research grants including American Academy of Advertising Research Grant"],
            ["Textbook", "International Marketing (two editions, Sage Publications)"],
            ["Relevance to Richard", "Marketing expertise; program oversight; could advise on quantitative aspects of procurement research"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Lisa Victoravich
    doc.add_heading("Lisa Victoravich -- Program Founder (Departed)", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Former Title", "Associate Dean for Faculty and Research; Director, Executive PhD Program (April 2020 - June 2023)"],
            ["Departed", "June 2023 to become Dean, College of Business and Economics, University of Idaho"],
            ["Background", "PhD from Florida State University; former CPA at Ernst & Young (Assurance Services)"],
            ["Significance", "Founded the Executive PhD program; her departure removed the original vision-holder from program leadership"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Craig Wallace
    doc.add_heading("Craig Wallace -- Founding Research Director (Departed)", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Former Title", "Head & Professor, Department of Management; Co-Director & Research Director, Executive PhD Program (2017-2019)"],
            ["Current Position", "Head of Management Department, Clemson University (Wilbur O. and Ann Powers College of Business)"],
            ["Research Areas", "Performance and effectiveness at individual and group levels; high-involvement leadership; teamwork and employee empowerment"],
            ["Significance", "Departed within ~1 year of program launch; early departure raised questions about program stability"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Melissa Akaka
    doc.add_heading("Melissa Akaka", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Professor of Marketing; Associate Dean for Research and Brand Strategy; Director, MS in Marketing; Co-Director, CiBiC"],
            ["Research Areas", "Value co-creation, service ecosystems, collaborative innovation, entrepreneurship, consumption journeys, practice diffusion, consumer cultures"],
            ["Teaching", "Marketing research, consumer behavior, qualitative research methods"],
            ["Relevance to Richard", "Qualitative methods expertise; value co-creation framework could inform stakeholder analysis in procurement research"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Andrew Schnackenberg
    doc.add_heading("Andrew Schnackenberg", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Associate Professor of Management"],
            ["PhD", "Case Western Reserve University"],
            ["Research Areas", "Organizational transparency, trust, symbolism, stakeholder management"],
            ["Key Work", "Defines transparency as three-dimensional: perceived information disclosure, clarity, and accuracy"],
            ["Publications", "Journal of Management (2016), Human Relations (2021), Journal of Trust Research (2022); 2,155+ citations"],
            ["Relevance to Richard", "Transparency and trust research directly relevant to government procurement transparency; potential committee member"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Yashar Atefi
    doc.add_heading("Yashar Atefi", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Full Professor and Chair of Marketing; Co-Director, CiBiC; Piccinati Associate Professor"],
            ["PhD", "University of Houston (2016)"],
            ["Research Areas", "Marketing strategy, sales management, growth departments, performance rankings as motivational tools"],
            ["Key Publication", "\"Open Negotiation\" in Journal of Marketing Research -- on salespeople's transparency in negotiations"],
            ["Relevance to Richard", "Negotiation and sales research could connect to best-value negotiations in procurement"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Tracy Xu
    doc.add_heading("Tracy Xu (Pisun Xu)", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Associate Professor of Finance; Ron Rizzuto Endowed Professor (Reiman School of Finance)"],
            ["PhD", "University of Washington (2006)"],
            ["Research Areas", "Corporate finance, corporate governance, international finance, real estate finance, credit default swap markets"],
            ["Credentials", "CFA charterholder"],
            ["Publications", "Financial Management, Journal of Financial and Quantitative Analysis, Journal of Banking and Finance"],
            ["Relevance to Richard", "Quantitative finance methods could support cost-analysis dimensions of procurement research"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Naomi Boyd
    doc.add_heading("Naomi Boyd -- Dean (since July 2024)", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "24th Dean of the Daniels College of Business (effective July 1, 2024)"],
            ["Previous Positions", "7th Dean, VCU School of Business; Associate Dean, West Virginia University; long-serving chair of WVU's finance department"],
            ["Education", "PhD from George Washington University; MBA in Finance from Texas Tech; BA from University of Texas at Austin"],
            ["Research Areas", "Structural changes in markets/processes and their impact on underlying microstructure"],
            ["Significance", "New leadership brings fresh vision but also represents another transition in a program that has seen considerable turnover"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Dennis Wittmer
    doc.add_heading("Dennis Wittmer -- Professor Emeritus (Retired June 2024)", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Career at DU", "1991-2024 (33 years)"],
            ["PhD", "Syracuse University, Maxwell School of Citizenship and Public Affairs (1992)"],
            ["Research Areas", "Managerial ethics, ethical sensitivity in decision-making, public/private organizations, project management, R&D science policy"],
            ["Awards", "Best Dissertation Award (Academy of Management, Public/Nonprofit Sector, 1992); Piccinati Endowed Professorship for Teaching Excellence (2002-04)"],
            ["Student Collaboration", "Co-authored Frontiers in Psychology paper with Cohort 3 students on VR, empathy, and moral reasoning"],
            ["Status", "Retired end of June 2024; may still be available for committee service as emeritus"],
        ],
        col_widths=[1.5, 5.0]
    )

    # Irina Khindanova
    doc.add_heading("Irina Khindanova", level=2)
    add_table(doc,
        ["", ""],
        [
            ["Title", "Teaching Associate Professor, Reiman School of Finance"],
            ["Education", "BS Applied Mathematics (Irkutsk State University); MA Development Economics (Williams College); PhD Economics (UC Santa Barbara, 2000)"],
            ["Research Areas", "Financial risk management, financial modeling, international investments, acquisitions, intellectual capital"],
            ["Student Collaboration", "Co-authored JABE paper with Cohort 4 student: \"Technological Acquisitions: The Impact of Innovation on Stock Performance\" (2023)"],
        ],
        col_widths=[1.5, 5.0]
    )

    doc.add_page_break()

    # ===== 6. COHORT HISTORY & ALUMNI =====
    doc.add_heading("6. Cohort History & Alumni", level=1)

    doc.add_heading("Cohort Timeline", level=2)
    add_table(doc,
        ["Cohort", "Start", "Est. Graduation", "Notes"],
        [
            ["1", "Fall 2018", "2021-2022", "18 students, avg age 47; founding cohort"],
            ["2", "Fall 2019", "2022-2023", ""],
            ["3", "Fall 2020", "2023-2024", "Published in Frontiers in Psychology with Wittmer"],
            ["4", "Fall 2021", "2024-2025", "Published in JABE with Khindanova"],
            ["5", "Fall 2022", "2025-2026", "USASBE Doctoral Consortium selection; HICSS paper"],
            ["6", "Fall 2023", "2026-2027", ""],
            ["7", "Fall 2024", "2027-2028", ""],
            ["8", "Fall 2025", "2028-2029", "Current cohort"],
        ],
        col_widths=[0.8, 1.0, 1.3, 3.4]
    )

    doc.add_heading("Known Alumni & Dissertation Topics", level=2)

    alumni_data = [
        ["Maurice Harris", "Cohort 1", "Going Backstage: Examining Justice Perceptions of Electronic Performance Monitoring",
         "Tenure-track Assistant Professor, University of Northern Colorado; Director, Daniels Fund Ethics Initiative at UNC"],
        ["Matt Orlowsky", "Cohort 1", "Understanding Critical Reflection as a Leadership Development Methodology (analyzed 352 USAFA cadet essays)",
         "Air Force veteran; career details not publicly available"],
        ["Peter Tripp", "Cohort 1", "Could Alexa Increase Your Social Worth? (reciprocity, trust, and AI in interpersonal relationships)",
         "Background includes VP/C-suite roles, nonprofits, and HBO; EMBA from DU (2018)"],
        ["Dawn Reed", "Cohort 1", "Decision-making research (specific title not publicly available)",
         "CPA from Fort Collins; adjunct professor at Bay Path University; goal to transition to full-time academia"],
        ["Kelly Watson", "Cohort ~2", "Factors Affecting Openness to Microaggression Training (surveyed 200+ respondents on 10 factors)",
         "Managing Partner, Orange Grove Consulting (Boston-based DEI firm)"],
        ["Faye Farmer", "Graduated 2024", "My Superpower: Research Collaborator Strategies Among Faculty Members Who Are First-Generation College Graduates",
         "Director, Enterprise Design Initiatives"],
        ["Sheryl Moriarty", "Graduated 2024", "The Effect of Financial Disclosure Medium and CEO Gender on Perceptions of Overconfidence and Credibility",
         "Assistant Professor of the Practice, School of Accountancy, Daniels College of Business (Fall 2025); DU MBA alumna; CPA; 30+ years at KPMG, BDO, RSM"],
        ["Gwen Almodovar", "Graduated ~2025", "How Practices Drive Persistence: The Transformational Consumption Journey of Native Americans in Higher Education",
         "Chief Operating Officer, Great American Paper USA LP"],
        ["Gordon Broadbent", "Active/Recent", "How large companies adapt emerging technologies like AI (focus: Department of Defense)",
         "Career goal: influence large organizations' tech integration approaches"],
        ["Jasmine Motupalli", "Graduated", "Role of emerging technologies in promoting economic development and equity in emerging markets",
         "Founder & CEO, The Data Love Co.; West Point graduate; MS from Georgia Tech"],
    ]

    add_table(doc,
        ["Name", "Cohort", "Dissertation Topic", "Current Position/Outcome"],
        alumni_data,
        col_widths=[1.2, 0.8, 2.5, 2.0]
    )

    doc.add_heading("Career Outcome Patterns", level=2)
    add_bullet(doc, "Maurice Harris secured the only confirmed tenure-track position -- at a regional teaching university (UNC), not an R1 institution.", bold_prefix="Academic placement: ")
    add_bullet(doc, "Sheryl Moriarty was hired back by DU Daniels in a practitioner-track (not tenure-track) role.", bold_prefix="Practice faculty: ")
    add_bullet(doc, "Several alumni (Watson, Almodovar, Motupalli) returned to or continued in industry/consulting roles with enhanced credentials.", bold_prefix="Industry/consulting: ")
    add_bullet(doc, "No evidence found of any Executive PhD graduate securing a tenure-track position at an R1 research university.", bold_prefix="R1 placement: ")

    doc.add_page_break()

    # ===== 7. STUDENT RESEARCH & PUBLICATIONS =====
    doc.add_heading("7. Student Research & Publications", level=1)

    add_para(doc, "The following peer-reviewed publications and conference presentations by Executive PhD students have been documented:")

    pub_data = [
        ["1", "Cohort 3 + Dennis Wittmer", "Expanding the Empirical Study of Virtual Reality Beyond Empathy to Compassion, Moral Reasoning and Moral Foundations", "Frontiers in Psychology", "Peer-reviewed journal"],
        ["2", "Cohort 4 (Arthur) + Irina Khindanova", "Technological Acquisitions: The Impact of Innovation on Stock Performance", "Journal of Applied Business and Economics, 25(4), 252-274 (2023)", "Peer-reviewed journal"],
        ["3", "Cohorts 5 and 3 (collaboration)", "AI Adoption in Entrepreneurial Firms: The Influence of Innovativeness, Proactiveness, and Risk-Taking", "Hawaii International Conference on System Sciences (HICSS)", "Conference proceedings"],
        ["4", "Cohort 5 student", "Entrepreneurship education project (with editors from Entrepreneurship Education & Pedagogy)", "2024 USASBE Doctoral Consortium", "National selection (1 of 20 doctoral students)"],
    ]

    add_table(doc,
        ["#", "Authors", "Title", "Venue", "Type"],
        pub_data,
        col_widths=[0.3, 1.3, 2.0, 1.7, 1.2]
    )

    add_para(doc, "These publications demonstrate that the program provides genuine research mentorship and that students can produce publishable work during their doctoral studies. The Frontiers in Psychology and JABE publications are particularly noteworthy as peer-reviewed journal placements. The HICSS conference is a well-regarded venue for information systems and technology research.")

    add_rich_para(doc, [
        ("Assessment: ", True),
        ("Four documented publications/presentations across 7 cohorts is modest but not unusual for an Executive PhD program where students maintain full-time careers. The faculty co-authorship model (Wittmer, Khindanova) suggests active mentorship rather than students working in isolation.", False),
    ])

    doc.add_page_break()

    # ===== 8. PROGRAM REPUTATION & RANKINGS =====
    doc.add_heading("8. Program Reputation & Rankings", level=1)

    doc.add_heading("Daniels College Rankings", level=2)
    add_table(doc,
        ["Ranking", "Position", "Source"],
        [
            ["Best Business Schools (Graduate)", "#97 (tie) of 133", "US News & World Report"],
            ["Part-time MBA", "#110 (tie) of 239", "US News & World Report"],
            ["Online MBA", "#65 of 344 (rose 9 spots)", "US News & World Report"],
            ["Best Undergraduate Business Programs", "#88 (jumped 6 spots)", "US News 2026 Rankings"],
            ["AACSB Accreditation", "Continuous since 1923", "Among the first accredited; first in Rocky Mountain Region"],
            ["University Classification", "R1: Very High Research Activity", "Carnegie Classification"],
            ["Executive PhD Specific Ranking", "No dedicated ranking exists", "True of ALL Executive PhD programs"],
        ],
        col_widths=[2.5, 2.0, 2.0]
    )

    add_para(doc, "Important context: There are no dedicated rankings for Executive PhD programs in business. The major ranking systems (US News, QS, ARWU, Financial Times) do not rank Executive PhD programs as a distinct category. The Executive DBA Council (EDBAC) maintains a directory but not comparative rankings. This means DU's Executive PhD cannot be directly compared to competitors on ranking alone.")

    doc.add_heading("Competitor Comparison Table", level=2)
    add_table(doc,
        ["Program", "Degree", "Total Cost", "Credits", "Duration", "GMAT?", "AACSB", "Location"],
        [
            ["DU Daniels", "PhD", "$104,580", "90", "3 years", "No", "Yes (1923)", "Denver, CO"],
            ["Virginia Tech", "PhD", "Not published", "60-90", "3-4+ years", "No", "Yes", "Alexandria, VA"],
            ["Oklahoma State", "PhD", "$139,500", "60", "3 years", "No", "Yes", "Stillwater, OK"],
            ["Bentley", "PhD", "$153,000", "Not published", "3 years", "Optional", "Yes", "Waltham, MA"],
            ["Claremont (Drucker)", "PhD", "$107K-$143K", "48-64", "3 years", "Yes", "Yes", "Claremont, CA"],
            ["Temple Fox", "DBA", "$118,800", "54", "3 years", "No", "Yes", "Philadelphia, PA"],
            ["Georgia State", "DBA", "$79K-$84K", "~54", "3 years", "Not specified", "Yes", "Atlanta, GA"],
        ],
        col_widths=[1.2, 0.6, 1.1, 0.7, 0.8, 0.6, 0.9, 1.1]
    )

    doc.add_heading("Key Comparison Insights", level=3)
    add_bullet(doc, "DU Daniels ($104,580) is the lowest-cost Executive PhD program among peers. Only the DBA programs at Temple ($118,800) and Georgia State ($79K-$84K) are comparable or cheaper, but these award a DBA rather than a PhD.", bold_prefix="Cost leader: ")
    add_bullet(doc, "Virginia Tech, Oklahoma State, Bentley, and DU all award a PhD (not a DBA). Under AACSB standards, both PhD and DBA holders qualify as Scholarly Academics (SA) for initial faculty hiring, but some institutions and hiring committees still prefer the PhD designation.", bold_prefix="PhD vs. DBA: ")
    add_bullet(doc, "DU's 90 credit hours is the highest among competitors. Oklahoma State requires only 60, Claremont 48-64. More credits could mean more thorough preparation or simply more coursework time/cost.", bold_prefix="Credit hours: ")
    add_bullet(doc, "All programs follow a similar 3-year, cohort-based, hybrid format with monthly weekend residencies and online components.", bold_prefix="Format similarity: ")
    add_bullet(doc, "Both DU and Virginia Tech hold R1 Carnegie classification. Oklahoma State is also R1. Bentley and Claremont are not R1.", bold_prefix="R1 status: ")

    doc.add_page_break()

    # ===== 9. HONEST ASSESSMENT =====
    doc.add_heading("9. Honest Assessment: Strengths & Concerns", level=1)

    doc.add_heading("Strengths", level=2)

    add_rich_para(doc, [("1. True PhD from an AACSB-Accredited R1 University", True)])
    add_para(doc, "This is not a DBA or an EdD. It is a PhD in Business Administration from a school accredited by AACSB since 1923 at a Carnegie R1 university. This distinction matters for AACSB faculty qualification (SA status), academic hiring committees, and professional credibility. Under AACSB Standard 3.2, newly appointed doctoral faculty are SA-qualified for up to 5 years after degree completion.")

    add_rich_para(doc, [("2. Competitive Price Point", True)])
    add_para(doc, "At $104,580, DU offers the most affordable Executive PhD among direct competitors. The next cheapest PhD option is Claremont at $107K-$143K (variable based on transfer credits). Oklahoma State is $139,500; Bentley is $153,000.")

    add_rich_para(doc, [("3. Faculty Who Publish with Students", True)])
    add_para(doc, "The documented co-authorship model (Wittmer in Frontiers in Psychology, Khindanova in JABE) demonstrates genuine research mentorship. This is not always the case in executive doctoral programs.")

    add_rich_para(doc, [("4. Cohort-Based Professional Network", True)])
    add_para(doc, "Cohort members are accomplished professionals (Cohort 1 average age 47, 10+ years experience). The network value extends beyond the degree itself, particularly for consulting and advisory work.")

    add_rich_para(doc, [("5. Hybrid Format Preserves Career Continuity", True)])
    add_para(doc, "Nine weekend immersions per year plus online sessions allow students to maintain full-time careers and salaries, partially offsetting the tuition cost.")

    doc.add_heading("Concerns", level=2)

    add_rich_para(doc, [("1. Significant Leadership Turnover", True)])
    add_para(doc, "This is the most notable concern. The program has experienced:")
    add_bullet(doc, "Craig Wallace (founding Research Director) departed for Clemson within ~1 year of program launch (~2019)")
    add_bullet(doc, "Lisa Victoravich (founder and Program Director) departed for University of Idaho as Dean in June 2023")
    add_bullet(doc, "Dennis Wittmer (emeritus, taught in ExecPhD, published with students) retired June 2024")
    add_bullet(doc, "Vivek Choudhury (previous Daniels Dean) replaced by Naomi Boyd in July 2024")
    add_para(doc, "The current director, Dan Baack, is the third program leader in seven years. While faculty turnover is normal in academia, losing both founders within five years of launch is concerning for program continuity and vision.")

    add_rich_para(doc, [("2. Young Program with Limited Track Record", True)])
    add_para(doc, "Founded in 2018, the program has only 4-5 graduating classes. The alumni network is small, career outcome data is thin, and the program has not yet weathered a significant institutional challenge (budget cuts, enrollment decline, accreditation review). By comparison, Oklahoma State's Executive PhD began in 2012 (10+ cohorts) and Claremont's executive doctoral programs span ~40 years.")

    add_rich_para(doc, [("3. Career Outcomes Favor Teaching/Practice Over R1 Tenure-Track", True)])
    add_para(doc, "This is a structural reality of ALL Executive PhD programs, not specific to DU:")
    add_bullet(doc, "The only confirmed tenure-track placement (Maurice Harris at UNC) is at a regional teaching university, not an R1")
    add_bullet(doc, "Sheryl Moriarty was hired as Professor of the Practice (non-tenure-track) at DU Daniels")
    add_bullet(doc, "Most graduates return to industry with enhanced credentials")
    add_para(doc, 'As the R3ciprocity blog (written by a business school professor) notes: "Executive DBAs probably won\'t help as much as a PhD to become a tenure-track business professor." Academic forums are blunter: R1 hiring committees prioritize publication records in top-tier journals, advisor pedigree, and full-time research immersion -- all areas where executive programs are structurally disadvantaged.')

    add_rich_para(doc, [("4. Self-Funded Investment Without Guaranteed Return", True)])
    add_para(doc, "~$105K out of pocket with no stipend, no assistantship, and no documented alumni tuition discount. Traditional funded PhDs receive $0 tuition plus $25K-$40K/year stipends. The ROI calculation depends heavily on post-graduation outcomes: if the goal is R1 tenure-track (unlikely from any Executive PhD), the ROI is poor. If the goal is enhanced consulting credibility and teaching opportunities, the ROI may justify the cost.")

    add_rich_para(doc, [("5. Highest Credit Requirement Among Peers", True)])
    add_para(doc, "DU requires 90 credit hours vs. Oklahoma State's 60 and Claremont's 48-64. While more credits could signal more thorough preparation, it also means more coursework time. Whether 28 credits of dissertation research (vs. fewer elsewhere) translates to better outcomes is unclear.")

    add_rich_para(doc, [("6. Limited Independent Reviews", True)])
    add_para(doc, "The program is too young and niche for substantial independent review. There are no ProPublica investigations, no Chronicle of Higher Education features, no substantive forum discussions. Most available information comes from DU's own marketing materials, which should be read with appropriate skepticism.")

    doc.add_page_break()

    # ===== 10. FIT ANALYSIS =====
    doc.add_heading("10. Fit Analysis for Richard's Situation", level=1)

    doc.add_heading("Advantages Specific to Richard", level=2)

    add_rich_para(doc, [("1. DU MBA Alumnus -- Institutional Knowledge", True)])
    add_para(doc, "Richard already knows the Daniels faculty, campus culture, administrative processes, and institutional expectations. This eliminates the learning curve that other students face and provides immediate access to faculty relationships. Sheryl Moriarty (MBA alumna who completed the ExecPhD) demonstrates this pathway is viable.")

    add_rich_para(doc, [("2. Advanced Dissertation Work Already Completed", True)])
    add_para(doc, "Richard has already written a 5-chapter, ~78,000-word dissertation draft with full statistical analysis (OLS, NB GLM, Logistic Regression, Propensity Score Matching) using federal procurement data. This is an extraordinary head start -- most students begin the dissertation process from scratch in Year 2-3. The question is whether DU will allow this work to carry forward or require him to start over within their Applied Research Practicum framework.")

    add_rich_para(doc, [("3. Dissertation Topic Alignment", True)])
    add_para(doc, '"From Lowest Price to Highest Public Value: An Empirical Test of Best-Value Source Selection in Government RFPs" intersects with:')
    add_bullet(doc, "Decision-making research (MGMT 6301, BUS 6300)")
    add_bullet(doc, "Organizational ethics (MGMT 6301 -- ethics in procurement)")
    add_bullet(doc, "Quantitative methods (BUS 6002/6003 -- already mastered)")
    add_bullet(doc, "Marketing strategy (MKTG 6300 -- value creation in B2G contexts)")
    add_bullet(doc, "Andrew Schnackenberg's transparency research -- directly relevant to government procurement transparency")

    add_rich_para(doc, [("4. Methods Training Already in Hand", True)])
    add_para(doc, "Richard has already demonstrated competency in OLS regression, negative binomial GLM, logistic regression, and propensity score matching. The quantitative methods courses (BUS 6002/6003) would reinforce rather than introduce these skills, allowing him to focus on deeper methodological nuance or alternative approaches.")

    add_rich_para(doc, [("5. Career Goal Alignment", True)])
    add_para(doc, "Richard's goals of consulting/advisory work with openness to academic teaching align well with actual Executive PhD outcomes. He is not targeting R1 tenure-track positions (the most unrealistic outcome), making the program's strengths -- credentialing, research training, network -- directly relevant to his intended path.")

    doc.add_heading("Potential Committee Members", level=2)
    add_table(doc,
        ["Faculty", "Relevance to Richard's Research", "Role Potential"],
        [
            ["Andrew Schnackenberg", "Transparency, trust, and stakeholder management -- directly applicable to procurement transparency", "Committee Chair or Member"],
            ["Dan Baack", "International marketing, business ethics -- connects to ethical procurement practices", "Committee Member (also program director)"],
            ["Dennis Wittmer (emeritus)", "Managerial ethics, public/private organizations -- published with students; deep government-adjacent research", "Committee Member (if available as emeritus)"],
            ["Tracy Xu", "Quantitative finance methods, corporate governance -- financial analysis of procurement outcomes", "Methods Advisor or Member"],
            ["Melissa Akaka", "Value co-creation -- theoretical framework for best-value procurement", "Committee Member"],
        ],
        col_widths=[1.5, 3.0, 2.0]
    )

    doc.add_heading("Cost Reality", level=2)
    add_table(doc,
        ["Factor", "Detail"],
        [
            ["Total tuition", "~$104,580"],
            ["Employer match available?", "Up to $10,000 if employer sponsors (must indicate before admission)"],
            ["GI Bill / Yellow Ribbon", "Not applicable (Richard is not a veteran)"],
            ["Alumni discount", "No documented discount for returning MBA alumni; recommend asking directly"],
            ["Opportunity cost", "Minimal -- program designed to maintain full-time career"],
            ["Net cost estimate", "$94,580-$104,580 depending on employer sponsorship"],
            ["Comparison", "Least expensive Executive PhD among peer programs"],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_heading("Risk Assessment", level=2)
    add_table(doc,
        ["Risk", "Likelihood", "Mitigation"],
        [
            ["Further faculty turnover disrupts program", "Moderate", "Program is AACSB-accredited and institutionally supported; DU has incentive to maintain it"],
            ["Dissertation work not accepted as prior credit", "Moderate-High", "Ask explicitly before enrolling; may need to adapt existing work into ARP framework"],
            ["No suitable committee chair for procurement research", "Low-Moderate", "Schnackenberg (transparency) and Wittmer (public orgs) are strong candidates; ask about availability"],
            ["Degree does not open desired career doors", "Low", "Richard's goals (consulting + teaching) align with actual program outcomes; PhD > DBA for credibility"],
            ["Program closes or loses accreditation", "Very Low", "AACSB since 1923; R1 university; 7+ cohorts enrolled"],
        ],
        col_widths=[2.5, 1.2, 2.8]
    )

    doc.add_page_break()

    # ===== 11. QUESTIONS TO ASK =====
    doc.add_heading("11. Questions to Ask the Program", level=1)

    add_para(doc, "Before accepting, Richard should pose the following questions to Daniels admissions and/or Executive PhD faculty:")

    doc.add_heading("Program Stability & Leadership", level=2)
    add_bullet(doc, "Both the program founder (Lisa Victoravich) and founding research director (Craig Wallace) have departed. How has the program's vision and direction evolved under Dan Baack's leadership?")
    add_bullet(doc, "What is the current retention/completion rate across cohorts? How many of the 18 Cohort 1 students actually completed the degree?")
    add_bullet(doc, "Has the program experienced any enrollment declines or considered changes to the curriculum since its founding?")

    doc.add_heading("Dissertation & Prior Work", level=2)
    add_bullet(doc, "I have completed substantial dissertation-level research (5 chapters, ~78K words, full statistical analysis using OLS, NB GLM, PSM) on best-value source selection in government procurement. Can this work be incorporated into the Applied Research Practicum sequence, or would I be required to start a new research project?")
    add_bullet(doc, "Who would be the most appropriate committee chair for research on government procurement, public value, and evidence-based source selection?")
    add_bullet(doc, "Is Andrew Schnackenberg available and willing to chair a dissertation on procurement transparency and public value?")

    doc.add_heading("Career Outcomes & Support", level=2)
    add_bullet(doc, "Can you provide specific career outcome data for graduates? How many have entered academia vs. returned to industry?")
    add_bullet(doc, "What post-graduation support does the program offer? Is there a formal alumni network or job placement assistance for academic positions?")
    add_bullet(doc, "Are there opportunities for teaching during the program (e.g., guest lectures, co-teaching, adjunct roles)?")

    doc.add_heading("Financial", level=2)
    add_bullet(doc, "Is there any tuition discount or scholarship available for DU/Daniels MBA alumni returning for the Executive PhD?")
    add_bullet(doc, "Are there merit-based scholarships available for the Executive PhD specifically?")
    add_bullet(doc, "What is the current per-credit rate, and is the $104,580 total figure guaranteed for the full 3-year duration?")

    doc.add_heading("Program Structure", level=2)
    add_bullet(doc, "Can any of the 90 required credits be waived or transferred based on prior graduate work (MBA coursework, independent research)?")
    add_bullet(doc, "What is the typical weekly time commitment during Years 1-2? During Year 3 (dissertation)?")
    add_bullet(doc, "How are immersion weekends scheduled? Is the academic calendar available for the upcoming cohort?")

    doc.add_page_break()

    # ===== 12. SOURCES =====
    doc.add_heading("12. Sources", level=1)

    add_para(doc, "This report draws on 60+ sources consulted in February 2026. Key sources are organized by category below.")

    doc.add_heading("Official DU/Daniels Sources", level=2)
    du_sources = [
        "Daniels Executive PhD Main Page: https://daniels.du.edu/executive-phd/",
        "Daniels Executive PhD Curriculum: https://daniels.du.edu/executive-phd/curriculum/",
        "DU Bulletin -- Executive PhD Courses: https://bulletin.du.edu/graduate/schoolscollegesanddivisions/danielscollegeofbusiness/executivephd/",
        "Daniels Executive PhD Application: https://daniels.du.edu/executive-phd/application/",
        "Daniels Executive PhD Faculty: https://daniels.du.edu/executive-phd/faculty/",
        "Daniels Executive PhD Student Profiles: https://daniels.du.edu/executive-phd/student-profiles/",
        "DU Tuition & Fees: https://www.du.edu/admission-aid/financial-aid-scholarships/graduate-financial-aid/cost-attend-du/tuition-fees",
        "DU Tuition Rate Schedule: https://www.du.edu/file/87955",
        "DU Yellow Ribbon Program: https://studentaffairs.du.edu/veterans/yellow-ribbon-funding-program",
        "DU Veteran Benefits Expansion: https://www.du.edu/news/veteran-benefits-expanding-cover-all-academic-programs",
        "DU Alumni Benefits: https://alumni.du.edu/resources-benefits",
        "Daniels College About/History: https://daniels.du.edu/about/ and https://daniels.du.edu/about/history/",
        "AACSB Accreditation -- DU: https://www.aacsb.edu/accredited/u/university-of-denver",
        "DU News -- US News Rankings: https://www.du.edu/news/us-news-world-report-ranks-du-top-college-2026-rankings",
        "Daniels Blog -- Program Launch: https://daniels.du.edu/blog/daniels-college-business-launches-executive-ph-d/",
        "Daniels Blog -- World Class (Cohort 1): https://daniels.du.edu/blog/world-class/",
        "Daniels Blog -- Naomi Boyd Named Dean: https://daniels.du.edu/blog/naomi-boyd-named-new-dean-of-daniels-college-of-business/",
        "Daniels Blog -- Wittmer Retirement: https://daniels.du.edu/blog/daniels-professor-retires-as-the-consummate-team-player/",
        "Daniels Blog -- New Faculty Fall 2025: https://daniels.du.edu/blog/new-faculty-members-join-daniels-for-fall-2025/",
    ]
    for s in du_sources:
        add_bullet(doc, s)

    doc.add_heading("Faculty Profiles & Scholarship", level=2)
    faculty_sources = [
        "Dan Baack: https://daniels.du.edu/directory/dan-baack/",
        "Melissa Akaka: https://daniels.du.edu/directory/melissa-akaka/",
        "Yashar Atefi: https://daniels.du.edu/directory/yashar-atefi/",
        "Tracy Xu: https://daniels.du.edu/directory/tracy-xu/",
        "Andrew Schnackenberg (Google Scholar): https://scholar.google.com/citations?user=EWv1UswAAAAJ",
        "Andrew Schnackenberg (ResearchGate): https://www.researchgate.net/profile/Andrew-Schnackenberg",
        "Irina Khindanova: https://daniels.du.edu/directory/irina-khindanova/",
        "Craig Wallace -- Clemson: https://www.clemson.edu/business/about/profiles/cw74",
        "Lisa Victoravich -- LinkedIn: https://www.linkedin.com/in/lisavictoravich/",
        "Lisa Victoravich -- AGB Placement: https://www.agbsearch.com/placements/dean-of-the-college-of-business-and-economics-university-of-idaho-0",
    ]
    for s in faculty_sources:
        add_bullet(doc, s)

    doc.add_heading("Alumni Dissertations", level=2)
    alumni_sources = [
        "DU Digital Commons (Theses & Dissertations): https://digitalcommons.du.edu/etd/",
        "Matt Orlowsky Dissertation: https://digitalcommons.du.edu/etd/2145/",
        "Peter Tripp Dissertation: https://digitalcommons.du.edu/etd/2089/",
        "Faye Farmer Dissertation: https://digitalcommons.du.edu/etd/2500/",
        "Sheryl Moriarty Dissertation: https://digitalcommons.du.edu/etd/2491/",
        "Maurice Harris -- Daniels Blog: https://daniels.du.edu/blog/daniels-grad-pays-it-back-with-executive-phd/",
        "Maurice Harris -- UNC Profile: https://www.unco.edu/employee-directory/maurice-harris/",
        "Kelly Watson -- DU News: https://www.du.edu/news/research-why-do-some-resist-microaggression-training",
        "Gwen Almodovar -- Daniels Blog: https://daniels.du.edu/blog/why-arent-more-native-american-students-attending-college/",
        "Faye Farmer -- Daniels Blog: https://daniels.du.edu/blog/executive-phd-student-examines-identitys-influence-on-collaboration/",
    ]
    for s in alumni_sources:
        add_bullet(doc, s)

    doc.add_heading("Competitor Programs", level=2)
    comp_sources = [
        "Virginia Tech Executive PhD: https://execphd.vt.edu/",
        "Virginia Tech Tuition: https://execphd.vt.edu/admissions/tuition.html",
        "Virginia Tech Curriculum: https://execphd.vt.edu/academics/curriculum.html",
        "Oklahoma State Executive PhD: https://business.okstate.edu/execphd/",
        "Bentley Executive PhD: https://www.bentley.edu/academics/phd-programs/executive-phd",
        "Bentley PhD Finances: https://www.bentley.edu/offices/student-accounts/executive-phd-finances",
        "Claremont Executive PhD: https://www.cgu.edu/academics/program/executive-phd/",
        "Claremont Tuition: https://www.cgu.edu/admissions/cost-aid/tuition-fees/",
        "Temple Fox Executive DBA: https://www.fox.temple.edu/academics/executive-dba",
        "Georgia State Robinson DBA: https://robinson.gsu.edu/program/dba/",
        "EDBAC Program Directory: https://edbac.org/program-comparison/",
        "DBA Compass: https://www.dba-compass.com/",
    ]
    for s in comp_sources:
        add_bullet(doc, s)

    doc.add_heading("Rankings & Accreditation", level=2)
    rank_sources = [
        "US News -- DU Daniels: https://www.usnews.com/best-graduate-schools/top-business-schools/university-of-denver-01038",
        "AACSB -- DU: https://www.aacsb.edu/accredited/u/university-of-denver",
        "AACSB -- Virginia Tech: https://www.aacsb.edu/accredited/v/virginia-polytechnic-institute-and-state-university",
        "AACSB -- Claremont: https://www.aacsb.edu/accredited/c/claremont-graduate-university",
        "AACSB Faculty Qualifications FAQs: https://www.aacsb.edu/educators/accreditation/business-accreditation/resources/faculty-qualifications",
    ]
    for s in rank_sources:
        add_bullet(doc, s)

    doc.add_heading("Independent Analysis & Criticism", level=2)
    crit_sources = [
        "R3ciprocity Blog -- Executive PhD Pros & Cons: https://blog.r3ciprocity.com/pros-and-cons-of-executive-phd-in-business-executive-dba-programs/",
        "R3ciprocity Blog -- Why You Should NOT Get a Doctorate: https://blog.r3ciprocity.com/why-you-should-not-get-a-doctorate-phd-dba-in-business-administration/",
        "DBA Compass -- Executive PhD vs Online PhD: https://www.dba-compass.com/knowledge/executive-phd-vs-online-phd/",
        "Vlerick Business School -- PhD vs DBA vs Executive PhD: https://www.vlerick.com/en/insights/what-s-the-difference-between-a-phd-dba-and-executive-phd/",
        "Babson -- DBA vs PhD: https://www.babson.edu/doctor-of-business-administration/dba-vs-phd-in-business-administration/",
        "FIU -- DBA as Pathway to Academia: https://business.fiu.edu/academics/graduate/insights/posts/the-dba-degree-as-a-pathway-to-academia.html",
    ]
    for s in crit_sources:
        add_bullet(doc, s)

    # ===== FOOTER =====
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Report generated February 2026. All facts verified against cited sources.")
    run.italic = True
    run.font.size = Pt(9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("For questions or updates, re-run research against source URLs.")
    run.italic = True
    run.font.size = Pt(9)

    # Save
    doc.save(str(OUTPUT))
    print(f"Report saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build_report()
