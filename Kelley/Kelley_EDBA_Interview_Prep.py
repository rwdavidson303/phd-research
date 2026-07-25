#!/usr/bin/env python3
"""Generate Kelley EDBA Interview Prep Word Document"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import subprocess

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x32, 0x6B)  # IU Crimson-ish navy
    return h

def add_key_point(text):
    p = doc.add_paragraph()
    run = p.add_run('► ')
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    run.bold = True
    p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KELLEY SCHOOL OF BUSINESS\nEXECUTIVE DBA PROGRAM')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Admissions Interview Preparation Brief')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x32, 0x6B)

doc.add_paragraph()
date_line = doc.add_paragraph()
date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_line.add_run('Prepared for Richard Davidson\nMarch 10, 2026')
run.font.size = Pt(14)

interviewer = doc.add_paragraph()
interviewer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = interviewer.add_run('\nInterviewer: Allison O\'Boyle\nEDBA Admissions & Academic Advisor')
run.font.size = Pt(13)
run.italic = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Program Overview & History',
    '2. Program Structure & Format',
    '3. Curriculum',
    '4. Faculty & Leadership',
    '5. Class Size & Student Profile',
    '6. Admissions Requirements',
    '7. Rankings & Reputation',
    '8. Key Differentiators',
    '9. Notable Kelley Alumni',
    '10. Interviewer Profile: Allison O\'Boyle',
    '11. Likely Interview Questions & How to Answer',
    '12. Questions YOU Should Ask',
    '13. How to Position DU PhD vs. Kelley DBA',
    '14. How to Present Your Research',
    '15. Red Flags to Avoid',
    '16. Your Key Advantages',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================================
# 1. PROGRAM OVERVIEW & HISTORY
# ============================================================
add_heading_styled('1. Program Overview & History', level=1)

doc.add_paragraph(
    'The Kelley Executive Doctor of Business Administration (EDBA) is a brand-new program. '
    'It was officially announced in January 2026 and will enroll its inaugural cohort in Fall 2026. '
    'There are no prior graduates or existing alumni of this specific program — you would be in the first class.'
)

add_heading_styled('Why Kelley Launched the EDBA', level=2)
doc.add_paragraph(
    'Kelley identified a gap between the MBA and advanced doctoral-level leadership development for senior executives. '
    'With a 137,000+ alumni network, Kelley received demand signals from senior leaders facing complex challenges: '
    'institutional changes, heightened stakeholder expectations, technological disruption, and geopolitical pressures.'
)

p = doc.add_paragraph()
run = p.add_run('Dean Pat Hopkins stated: ')
run.bold = True
p.add_run('"[The program] builds on this tradition by elevating strategic leadership development through '
          'the integration of advanced business theory and practical application."')

doc.add_paragraph(
    'The program is delivered through Kelley Direct Programs, which also runs the #1 ranked online MBA program '
    '(U.S. News, five consecutive years), giving the EDBA immediate infrastructure credibility.'
)

# ============================================================
# 2. PROGRAM STRUCTURE & FORMAT
# ============================================================
doc.add_page_break()
add_heading_styled('2. Program Structure & Format', level=1)

add_table(
    ['Detail', 'Information'],
    [
        ['Duration', '2 years (with MBA transfer credits) or 3 years (without prior MBA)'],
        ['Total Credits', '69 credit hours (no MBA) or 42 credit hours (with MBA, up to 27 transferable)'],
        ['Modality', 'Hybrid — primarily online with required in-person residencies'],
        ['Live Sessions', 'One night per week, ~75 minutes, 6:00–9:00 PM Eastern Time (recorded)'],
        ['Weekly Commitment', 'Approximately 15–20 hours per week'],
        ['Tuition', '$2,300 per credit hour'],
        ['Total Cost (with MBA)', '~$96,600 plus residency fees'],
        ['Total Cost (without MBA)', '~$142,440 (MBA credits at MBA rate, remainder at EDBA rate)'],
        ['Residency Fees', '$1,000 per residency (additional)'],
        ['Concentration', 'Strategy and Leadership'],
    ]
)

add_heading_styled('Residency Details', level=2)

doc.add_paragraph('For students WITHOUT a prior MBA — 4 total residencies:', style='List Bullet')
doc.add_paragraph(
    'First Year: "Kelley On Campus" — Intensive introduction on IU Bloomington campus',
    style='List Bullet 2'
)
doc.add_paragraph(
    'Second Year: "Kelley On Location" — Choice of topics/formats; Bloomington or major U.S. city',
    style='List Bullet 2'
)
doc.add_paragraph(
    'EDBA Special Topics I (3 credits) — Executive presence, persuasion, data-driven communication. '
    'First cohort dates: November 5–9, 2026',
    style='List Bullet 2'
)
doc.add_paragraph(
    'EDBA Special Topics II (3 credits) — Thesis development: refining research questions, methods, '
    'identifying faculty advisors. Summer between year one and two.',
    style='List Bullet 2'
)
doc.add_paragraph('For students WITH a prior MBA — 2 doctoral residencies (items 3 and 4 above)', style='List Bullet')

# ============================================================
# 3. CURRICULUM
# ============================================================
doc.add_page_break()
add_heading_styled('3. Curriculum', level=1)

doc.add_paragraph('The 42 doctoral credit hours break into five sections:')

add_heading_styled('A. Strategy & Leadership Courses — 18 Credit Hours', level=2)
doc.add_paragraph(
    'Six courses focused on strategy and leadership. Faculty expertise includes: global strategy, '
    'corporate innovation, leadership, entrepreneurship, international business, consulting, negotiations, '
    'change management, organizational design, crisis management, and turnaround management.'
)

add_heading_styled('B. Statistics & Research Methods — 9 Credit Hours', level=2)
doc.add_paragraph(
    'Two statistics and research methods courses plus one thesis development course. '
    'Covers both quantitative and qualitative research methodologies.'
)

add_heading_styled('C. Thesis / Applied Research Project — 9 Credit Hours', level=2)
doc.add_paragraph(
    'Required applied research thesis with one-on-one faculty advising. Full research lifecycle: '
    'ideation, problem framing, data collection, analysis, writing, and committee presentation. '
    'Students have significant latitude in choosing topics — thesis should be tailored to the '
    'student\'s own organization or career context.'
)

add_heading_styled('D. EDBA Residencies — 6 Credit Hours', level=2)
doc.add_paragraph('Two in-person residencies at 3 credits each (described above).')

add_heading_styled('E. MBA Core Coursework — 27 Credit Hours (without MBA only)', level=2)
doc.add_paragraph('Foundational MBA-level courses charged at the MBA tuition rate.')

# ============================================================
# 4. FACULTY & LEADERSHIP
# ============================================================
doc.add_page_break()
add_heading_styled('4. Faculty & Leadership', level=1)

doc.add_paragraph(
    'The EDBA is taught by the same full-time, award-winning faculty who teach in the '
    'Kelley Direct MBA program (the #1 online MBA).'
)

add_heading_styled('Erik Gonzalez-Mule, Ph.D. — Associate Chair, Kelley Direct Programs', level=2)
doc.add_paragraph(
    'Holds the Randall L. Tobias Chair in Leadership. Chairperson, Department of Management and '
    'Entrepreneurship. Also holds the Eveleigh Professor of Business Leadership title. '
    'Ph.D. in Business Administration, University of Iowa (Presidential Research Fellow). '
    'Research: team composition and processes, stress, individual differences, employee health. '
    'Published in Journal of Applied Psychology, Journal of Management. '
    'Named 2024 Poets&Quants Best 40-Under-40 MBA Professor. '
    'Research featured in The Wall Street Journal, CBS News, Forbes.'
)

add_heading_styled('Patrick E. Hopkins, Ph.D. — Dean, Kelley School of Business', level=2)
doc.add_paragraph(
    'Effective March 17, 2025. Kelley faculty since 1995; previously Vice Dean. '
    'Ph.D. in Accounting, University of Texas at Austin (minors in psychology and statistics). '
    'CPA; won the Elijah Watt Sells Award (1987). Research: professional judgment, decision-making.'
)

add_heading_styled('Other Senior Leadership', level=2)
doc.add_paragraph('Dan Li — Executive Associate Dean; L. Leslie Waters Chair in International Business')
doc.add_paragraph('Stephanie Lu Wang — Samuel & Pauline Glaubinger Professor; Associate Chair, Management & Entrepreneurship')
doc.add_paragraph('Will Geoghegan — Kelley Direct Chair, Clinical Professor')
doc.add_paragraph('Miki Pike Hamstra — Executive Director of Kelley Direct')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run('The Management and Entrepreneurship Department alone has 50+ faculty members.')

# ============================================================
# 5. CLASS SIZE & STUDENT PROFILE
# ============================================================
doc.add_page_break()
add_heading_styled('5. Class Size & Student Profile', level=1)

add_table(
    ['Detail', 'Information'],
    [
        ['Target first cohort', '~25 students'],
        ['Maximum per admissions cycle', 'Capped at 50'],
        ['Target experience level', '10+ years of management experience'],
        ['Target profile', 'Mid-to-late career professionals; senior executives'],
        ['Geographic focus', 'Primarily U.S.-based (evening ET classes), international welcome'],
        ['Industries', 'Diverse (program emphasizes cross-industry peer networking)'],
    ]
)

add_heading_styled('Post-Graduation Career Paths', level=2)
doc.add_paragraph('Higher leadership roles within current organizations', style='List Bullet')
doc.add_paragraph('Academic careers (teaching, research)', style='List Bullet')
doc.add_paragraph('Independent consultancy', style='List Bullet')
doc.add_paragraph('Board positions', style='List Bullet')
doc.add_paragraph('Strategic advisory roles', style='List Bullet')

# ============================================================
# 6. ADMISSIONS REQUIREMENTS
# ============================================================
add_heading_styled('6. Admissions Requirements', level=1)

add_table(
    ['Requirement', 'Details'],
    [
        ['GMAT/GRE', 'Optional (not required)'],
        ['Work Experience', '10+ years of management experience (target)'],
        ['Prior Degree', 'MBA preferred but not required; those without complete extra 27 credits'],
        ['Application Fee', 'Phase 1 is FREE'],
        ['Priority Deadline', 'May 1, 2026 (for Fall 2026 cohort)'],
        ['Admissions Type', 'Rolling admission — submit early'],
    ]
)

add_heading_styled('Application Process', level=2)
doc.add_paragraph(
    'Phase 1: ~5 minutes; answer a few questions and upload resume (free). '
    'If a strong fit, contacted within 2 weeks. '
    'Phase 2: 30-minute interview with admissions team member. '
    'The FAQ states: "No need to be nervous — we love getting to know our applicants!"'
)

# ============================================================
# 7. RANKINGS & REPUTATION
# ============================================================
doc.add_page_break()
add_heading_styled('7. Rankings & Reputation', level=1)

add_table(
    ['Category', 'Ranking'],
    [
        ['Online MBA', '#1 — U.S. News, 5 consecutive years (same Kelley Direct platform as EDBA)'],
        ['Full-Time MBA', '#22 (tie) — U.S. News 2026'],
        ['Part-Time MBA', '#16 — U.S. News 2026'],
        ['Undergraduate Business', '#8 overall, #4 among publics — U.S. News 2026'],
        ['Research (UT-Dallas)', '#10 worldwide, #9 North America, #2 among publics'],
        ['Management Research (TAMUGA)', '#4 in U.S., #1 in Big Ten (tied with Michigan)'],
    ]
)

doc.add_paragraph()
add_key_point(
    'There are no established DBA-specific rankings. The program\'s credibility rests on '
    'AACSB accreditation, the #1 online MBA platform, and top-10 worldwide research productivity.'
)

# ============================================================
# 8. KEY DIFFERENTIATORS
# ============================================================
add_heading_styled('8. Key Differentiators', level=1)

differentiators = [
    ('Built on #1 Online MBA Platform', 'Same Kelley Direct infrastructure — proven technology, pedagogy, and support.'),
    ('Strategy & Leadership Focus', 'Single concentration provides depth over breadth. Not a broad survey.'),
    ('Scholar-Practitioner Model', 'Trains executives to be "critical consumers of information and data" who "question and improve organizations."'),
    ('Applied Thesis (Not Just Capstone)', 'Full applied research thesis — 9 dedicated credit hours with one-on-one faculty advising.'),
    ('Top-10 Research Faculty', 'UT-Dallas top 10 worldwide research. Students learn from actively publishing scholars.'),
    ('Lifetime Career Coaching', 'Kelley Direct Professional Advancement (KDPA) — lifetime access to one-on-one coaching.'),
    ('Flexible for Executives', 'One 75-min live class/week, recorded sessions, 15-20 hrs/week, 2-4 days residency/year.'),
    ('Cost Competitive', '~$96K–$142K. Competitive with Temple Fox, less than many European programs.'),
    ('Massive Alumni Network', '137,000+ Kelley alumni across 105 countries.'),
    ('AACSB Accredited', 'Full AACSB accreditation at the institutional level.'),
]

for title, desc in differentiators:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 9. NOTABLE KELLEY ALUMNI
# ============================================================
doc.add_page_break()
add_heading_styled('9. Notable Kelley Alumni', level=1)

p = doc.add_paragraph()
run = p.add_run('Important: The EDBA has NO alumni yet — the first cohort starts Fall 2026. ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('However, the broader Kelley network includes:')

alumni = [
    ('Mark Cuban', 'Billionaire entrepreneur, Shark Tank (BS Management \'81)'),
    ('David Solomon', 'CEO, Goldman Sachs (MBA \'92)'),
    ('John T. Chambers', 'Former CEO & Executive Chairman, Cisco Systems (MBA \'75)'),
    ('Alicia Boler Davis', 'CEO, Alto Pharmacy; former GM executive (MBA)'),
    ('Derica Rice', 'Former EVP, CVS Health; Board at Disney, Target, BMS (MBA \'90)'),
    ('Anton Vincent', 'President, Mars Wrigley North America (MBA \'93)'),
    ('Patrick Decker', 'President & CEO, Xylem'),
]
for name, desc in alumni:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name} — ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 10. INTERVIEWER PROFILE
# ============================================================
doc.add_page_break()
add_heading_styled('10. Interviewer Profile: Allison O\'Boyle', level=1)

p = doc.add_paragraph()
run = p.add_run('NOTE: ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('The interviewer\'s name is likely "Allison O\'Boyle" (not "Boyle"). Her IU email is asoboyle@iu.edu.')

doc.add_paragraph()
add_table(
    ['Detail', 'Information'],
    [
        ['Full Name', 'Allison O\'Boyle'],
        ['Title', 'EDBA Admissions and Academic Advisor'],
        ['Email', 'asoboyle@iu.edu'],
        ['Department', 'Kelley Direct Programs'],
        ['Location', '1275 E. Tenth Street, Suite 2100, Bloomington, IN 47405'],
        ['Role in Interview', 'She is the ONLY person listed specifically for EDBA admissions'],
    ]
)

add_heading_styled('What We Know', level=2)
doc.add_paragraph(
    'Allison holds a dual role as both EDBA Admissions contact and Academic Advisor. '
    'She is the dedicated point person for the EDBA program and the only person listed specifically '
    'for EDBA admissions on the Kelley Direct contact page.'
)

add_heading_styled('Reports To / Works With', level=2)
doc.add_paragraph('Abigail Gschwend — Director of Admissions and Financial Aid (likely her boss)', style='List Bullet')
doc.add_paragraph('Courtney Daily — Associate Director of Enrollment Management', style='List Bullet')
doc.add_paragraph('Erik Gonzalez-Mule — Associate Chair, Kelley Direct (EDBA academic lead)', style='List Bullet')
doc.add_paragraph('Miki Pike Hamstra — Executive Director of Kelley Direct', style='List Bullet')

add_heading_styled('Possible Family Connection', level=2)
doc.add_paragraph(
    'There is a notable faculty member at Kelley named Ernest O\'Boyle (eoboyle@iu.edu), '
    'who is the Dale M. Coleman Chair of Management and Professor in the Department of Management '
    'and Entrepreneurship. Ph.D. from Virginia Commonwealth University. Research: individual differences, '
    'Dark Triad personality, research methods. The shared last name could indicate a family relationship, '
    'though this is unconfirmed. If related, she likely has deep familiarity with academic research culture.'
)

add_heading_styled('What to Expect from Her', level=2)
doc.add_paragraph('Conversational, not adversarial — the FAQ literally says "we love getting to know our applicants"', style='List Bullet')
doc.add_paragraph('Fit-focused — she wants to understand your trajectory and why this program, why now', style='List Bullet')
doc.add_paragraph('Program-selling too — as recruiter for a brand-new program (inaugural cohort), she wants strong candidates', style='List Bullet')
doc.add_paragraph('Research-aware — as academic advisor for a doctoral program, she appreciates hearing about research interests', style='List Bullet')
doc.add_paragraph('30 minutes — be concise and prepared with your narratives', style='List Bullet')

# ============================================================
# 11. LIKELY INTERVIEW QUESTIONS & HOW TO ANSWER
# ============================================================
doc.add_page_break()
add_heading_styled('11. Likely Interview Questions & How to Answer', level=1)

add_heading_styled('Motivation & "Why" Questions', level=2)

questions_motivation = [
    ('"Why do you want to pursue a DBA/EDBA?"',
     'Lead with your vision for impact. Talk about how you want to bridge rigorous research and practical '
     'leadership in government procurement. Emphasize the scholar-practitioner identity.'),
    ('"Why now? What makes this the right time?"',
     'Connect your career trajectory to a natural inflection point. You have the experience and the research '
     'question already — now you need the doctoral framework to execute it rigorously.'),
    ('"Why a DBA rather than a PhD?"',
     'The DBA\'s emphasis on applied research and practitioner impact aligns better with your goal of '
     'producing research that directly changes how government agencies award contracts.'),
    ('"What do you hope to achieve that you cannot without this degree?"',
     'Be specific: doctoral-level research credibility, access to faculty mentorship, a rigorous thesis '
     'that can influence procurement policy. Avoid vague prestige answers.'),
]
for q, a in questions_motivation:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(f'How to answer: {a}')
    doc.add_paragraph()

add_heading_styled('Research Interest Questions', level=2)

questions_research = [
    ('"What are your research interests?"',
     'Your topic is gold: "From Lowest Price to Highest Public Value." Lead with the $700B+ problem, '
     'then your specific question, then mention you already have data and preliminary results.'),
    ('"Describe a business problem you want to study rigorously."',
     'Government awards contracts worth hundreds of billions based on evaluation methods that have never '
     'been empirically tested at scale. You are filling that gap.'),
    ('"How would you apply your research?"',
     'Direct policy implications: if best-value produces measurably better outcomes, procurement leaders '
     'should use it more aggressively. Your findings change how practitioners operate.'),
]
for q, a in questions_research:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(f'How to answer: {a}')
    doc.add_paragraph()

add_heading_styled('Professional Background Questions', level=2)
doc.add_paragraph('"Walk us through your career and how it led you here."', style='List Bullet')
doc.add_paragraph('"What is the most significant leadership challenge you have faced?"', style='List Bullet')
doc.add_paragraph('"Describe a time you used data or evidence to make an important decision."', style='List Bullet')

add_heading_styled('Capacity & Commitment Questions', level=2)
doc.add_paragraph('"How do you plan to balance the program with professional/personal responsibilities?"', style='List Bullet')
doc.add_paragraph('"What does your support system look like?"', style='List Bullet')

add_heading_styled('Kelley-Specific Questions', level=2)

kelley_specific = [
    ('"What does being a scholar-practitioner mean to you?"',
     'This is Kelley\'s core concept. Your answer: bridging rigorous research and practical leadership. '
     'You don\'t just want to lead — you want to lead with evidence.'),
    ('"How does strategy and leadership relate to your research?"',
     'Source selection methodology IS a strategic organizational decision. When a contracting officer '
     'chooses an evaluation method, they make a strategic choice that determines outcomes for taxpayer dollars.'),
    ('"What would you contribute to the cohort?"',
     'Government/public sector is likely underrepresented in a business school cohort. You bring a unique lens, '
     'plus you already have technical research skills (Python, statistical modeling, data pipelines).'),
    ('"What excites you about being part of the inaugural cohort?"',
     'Frame as: you are the kind of person who builds things from the ground up. You want to help shape '
     'what this program becomes.'),
]
for q, a in kelley_specific:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(f'How to answer: {a}')
    doc.add_paragraph()

# ============================================================
# 12. QUESTIONS YOU SHOULD ASK
# ============================================================
doc.add_page_break()
add_heading_styled('12. Questions YOU Should Ask', level=1)

doc.add_paragraph(
    'These demonstrate genuine research, intellectual curiosity, and strategic thinking. '
    'Pick 3–4 that feel most natural to you.'
)

questions_to_ask = [
    ('About Research & Thesis',
     [
         'How are students matched with faculty advisors, and how early does that process begin?',
         'My research interest is in government procurement — specifically, whether best-value source selection '
         'produces better outcomes. How does the program support students whose research sits at the intersection '
         'of public policy and organizational strategy?',
         'What does the thesis defense process look like? Traditional dissertation defense or more practitioner-oriented?',
     ]),
    ('About the Cohort Experience',
     [
         'With the inaugural cohort around 25 students, what professional backgrounds and industries are you '
         'seeing in the applicant pool?',
         'How are the in-residence experiences structured? Can you walk me through what those weeks look like?',
     ]),
    ('About Faculty & Intellectual Life',
     [
         'Are there faculty members whose work touches on organizational decision-making, procurement strategy, '
         'or public sector management?',
         'Will EDBA students have access to research seminars, working papers, or collaboration opportunities '
         'with PhD students?',
     ]),
    ('About Career & Post-Graduation',
     [
         'What outcomes is the program hoping to see from graduates? Enhanced practitioners, or also support '
         'for those who want to teach or publish?',
         'How does the program envision the EDBA alumni network developing as you build from the ground up?',
     ]),
    ('One Bold Question',
     [
         'As a new program, what is your vision for where the Kelley EDBA will be positioned in five years '
         'relative to established executive doctoral programs?',
     ]),
]

for category, questions in questions_to_ask:
    add_heading_styled(category, level=2)
    for q in questions:
        doc.add_paragraph(f'"{q}"', style='List Bullet')

# ============================================================
# 13. HOW TO POSITION DU PhD vs. KELLEY DBA
# ============================================================
doc.add_page_break()
add_heading_styled('13. How to Position DU PhD vs. Kelley DBA', level=1)

p = doc.add_paragraph()
run = p.add_run('This is the most delicate part of the interview. Be honest but strategic.')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

add_heading_styled('Recommended Language', level=2)
doc.add_paragraph(
    '"I have been accepted to a PhD program at the University of Denver\'s Daniels College of Business '
    'for Fall 2026, and I have been doing extensive preparation work on my research, including data collection '
    'and preliminary analysis. When I learned about Kelley\'s new Executive DBA program, I was drawn to it '
    'because [specific reasons]. I am taking this decision seriously because I want to choose the program '
    'that is the best fit for my research goals and professional trajectory."'
)

add_heading_styled('What to Emphasize', level=2)
doc.add_paragraph(
    'The DBA format may be a better fit for your career stage — applied research and scholar-practitioner model '
    'align with producing research that directly improves government contracting practices.',
    style='List Bullet'
)
doc.add_paragraph(
    'The hybrid/executive format respects your professional commitments.',
    style='List Bullet'
)
doc.add_paragraph(
    'Kelley\'s brand and network: R1 institution, #1 online MBA, top-10 research.',
    style='List Bullet'
)
doc.add_paragraph(
    'Strategy and leadership concentration maps to how government agencies make strategic sourcing decisions.',
    style='List Bullet'
)

add_heading_styled('The Strongest Frame', level=2)
doc.add_paragraph(
    '"Both programs would allow me to pursue my research, but the approaches differ. The PhD route is more '
    'traditionally academic, and the DBA route emphasizes applied impact and practitioner relevance. I am '
    'genuinely evaluating which approach will produce research that most effectively changes how government '
    'agencies award contracts — and which learning environment will push me the hardest. That is why I '
    'wanted to have this conversation."'
)

add_heading_styled('What NOT to Say', level=2)
doc.add_paragraph('Never disparage DU or Daniels.', style='List Bullet')
doc.add_paragraph('Don\'t say EDBA is a "backup" or "easier path."', style='List Bullet')
doc.add_paragraph('Don\'t frame it as "I got in somewhere else, but I\'d rather go here."', style='List Bullet')
doc.add_paragraph('Don\'t say "I\'m just exploring my options" — sounds uncommitted.', style='List Bullet')

# ============================================================
# 14. HOW TO PRESENT YOUR RESEARCH
# ============================================================
doc.add_page_break()
add_heading_styled('14. How to Present Your Research', level=1)

add_heading_styled('Lead with the Problem, Not the Method', level=2)
doc.add_paragraph(
    '"The U.S. federal government spends over $700 billion annually in contracts. The default approach — '
    'awarding to the lowest bidder — often produces poor outcomes: cost overruns, protests, and failed '
    'performance. Best-value source selection evaluates both price and non-price factors, but we lack '
    'rigorous empirical evidence on whether it actually delivers better results. My research fills that gap."'
)

add_heading_styled('Connect to Strategy & Leadership', level=2)
doc.add_paragraph(
    '"This is fundamentally a strategic decision-making question. When a contracting officer chooses an '
    'evaluation method, they make a strategic choice that determines whether the government gets the best '
    'possible outcome for taxpayer dollars. My research applies evidence-based analysis to one of the largest '
    'purchasing decisions any organization makes."'
)

add_heading_styled('Emphasize Applied/Practitioner Impact', level=2)
doc.add_paragraph(
    '"This research has direct implications for policy. If best-value source selection produces measurably '
    'better outcomes — fewer protests, better contractor performance, lower cost overruns — then procurement '
    'leaders should use it more aggressively. If it doesn\'t, we need to understand why. Either way, the '
    'findings change how practitioners operate."'
)

add_heading_styled('Show the Data Advantage', level=2)
doc.add_paragraph(
    '"I have already been working with USAspending.gov API data, FPDS contract records, and GAO bid protest data. '
    'I\'ve built data pipelines and run preliminary analyses including OLS regression, negative binomial models, '
    'logistic regression, and propensity score matching. I am not starting from scratch — I am bringing a '
    'working research program to the program."'
)

add_heading_styled('Frame as Universally Relevant', level=2)
doc.add_paragraph(
    '"While my specific context is government procurement, the underlying question — does paying for quality '
    'produce better outcomes than minimizing upfront cost? — applies to any organization that makes purchasing '
    'decisions. It is a question about organizational strategy and resource allocation."'
)

# ============================================================
# 15. RED FLAGS TO AVOID
# ============================================================
doc.add_page_break()
add_heading_styled('15. Red Flags to Avoid', level=1)

red_flags = [
    ('Don\'t appear to want the title without the work',
     'Show genuine intellectual curiosity about the research process, not just the credential.'),
    ('Don\'t be vague about research',
     'You have a specific, well-developed question with data already in hand. Use that advantage.'),
    ('Don\'t ask questions answered on the website',
     'Don\'t ask "How many credits?" or "Is it online?" — these signal you haven\'t done homework.'),
    ('Don\'t underestimate the rigor',
     'Avoid any language suggesting the DBA is easier than a PhD. The workload is substantial.'),
    ('Don\'t be dismissive of theory',
     'Say: "I want to test theory against real-world data to see what actually works."'),
    ('Don\'t badmouth your current situation',
     'Frame your pursuit as moving toward something, not running from something.'),
    ('Don\'t monologue',
     '30 minutes. Be concise. Answer directly, elaborate briefly. Leave room for conversation.'),
    ('Do show personality',
     'The FAQ says "we love getting to know applicants." Be professional but authentic.'),
]

for title, desc in red_flags:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    p.add_run(desc)

# ============================================================
# 16. YOUR KEY ADVANTAGES
# ============================================================
add_heading_styled('16. Your Key Advantages Going In', level=1)

advantages = [
    'You have a fully developed research question with a clear title, methodology, and data pipeline already built. Most DBA applicants walk in with vague ideas.',
    'Government procurement is a massive, understudied domain with direct policy implications — exactly the kind of applied research a DBA thesis should tackle.',
    'Your topic maps cleanly to "strategy and leadership" — source selection methodology is a strategic organizational decision.',
    'Public sector perspective diversifies the cohort. Business school cohorts skew private sector.',
    'You are technologically capable — Python, Jupyter, APIs, statistical modeling signals research readiness.',
    'Being in the inaugural cohort is something you can frame as exciting — you build things from the ground up.',
]

for i, adv in enumerate(advantages, 1):
    add_key_point(f'{adv}')

# ============================================================
# QUICK REFERENCE CARD (LAST PAGE)
# ============================================================
doc.add_page_break()
add_heading_styled('QUICK REFERENCE CARD — Interview Day', level=1)

p = doc.add_paragraph()
run = p.add_run('Interviewer: ')
run.bold = True
p.add_run('Allison O\'Boyle, EDBA Admissions & Academic Advisor')

p = doc.add_paragraph()
run = p.add_run('Format: ')
run.bold = True
p.add_run('30-minute interview (conversational)')

p = doc.add_paragraph()
run = p.add_run('Program: ')
run.bold = True
p.add_run('Kelley Executive DBA — Strategy & Leadership concentration')

p = doc.add_paragraph()
run = p.add_run('Key Fact: ')
run.bold = True
p.add_run('Inaugural cohort (Fall 2026) — you\'d be in the FIRST class')

p = doc.add_paragraph()
run = p.add_run('Platform: ')
run.bold = True
p.add_run('Kelley Direct — same as #1 ranked online MBA (5 consecutive years)')

p = doc.add_paragraph()
run = p.add_run('Academic Lead: ')
run.bold = True
p.add_run('Erik Gonzalez-Mule, Tobias Chair in Leadership')

p = doc.add_paragraph()
run = p.add_run('Dean: ')
run.bold = True
p.add_run('Pat Hopkins (since March 2025)')

doc.add_paragraph()
add_heading_styled('Your 30-Second Pitch', level=2)
doc.add_paragraph(
    '"I\'m a senior professional in government contracting with a passion for evidence-based decision-making. '
    'I\'ve spent my career seeing how procurement evaluation methods — particularly lowest-price versus '
    'best-value — drive dramatically different outcomes for taxpayers. I\'ve already started building the '
    'research to test this empirically, using federal contract data. The Kelley EDBA\'s focus on strategy '
    'and leadership, combined with the applied thesis model and world-class faculty, is exactly the doctoral '
    'framework I need to turn this research into actionable policy insights. And the opportunity to be part '
    'of the inaugural cohort is genuinely exciting to me."'
)

doc.add_paragraph()
add_heading_styled('Remember', level=2)
doc.add_paragraph('Be yourself — they want to know YOU, not a performance.', style='List Bullet')
doc.add_paragraph('You are also evaluating THEM — this is a mutual fit conversation.', style='List Bullet')
doc.add_paragraph('Your research preparation puts you ahead of most applicants.', style='List Bullet')
doc.add_paragraph('Ask your questions — it shows engagement and seriousness.', style='List Bullet')

# Save
output_path = '/Users/richardwdavidson/Desktop/Kelley_EDBA_Interview_Prep.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
