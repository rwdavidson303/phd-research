#!/usr/bin/env python3
"""Generate Kelley EDBA Interview Prep Word Document — Enhanced with Resume"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x32, 0x6B)
    return h

def add_key_point(text):
    p = doc.add_paragraph()
    run = p.add_run('>> ')
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    run.bold = True
    p.add_run(text)
    return p

def add_qa(question, answer):
    """Add a formatted question/answer pair"""
    p = doc.add_paragraph()
    run = p.add_run(question)
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    run2 = p2.add_run('YOUR ANSWER: ')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    run2.font.size = Pt(10)
    p2.add_run(answer).font.size = Pt(10)
    doc.add_paragraph()

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('KELLEY SCHOOL OF BUSINESS\nEXECUTIVE DBA PROGRAM')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Admissions Interview Preparation Brief')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x00, 0x32, 0x6B)

doc.add_paragraph()
date_line = doc.add_paragraph()
date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_line.add_run('Prepared for Richard Davidson, MBA\nMarch 10, 2026')
run.font.size = Pt(14)

interviewer = doc.add_paragraph()
interviewer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = interviewer.add_run('\nInterviewer: Allison O\'Boyle\nEDBA Admissions & Academic Advisor')
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run('PERSONALIZED EDITION — Answers tailored to your resume and experience')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Program Overview & History',
    '2. Program Structure & Format',
    '3. Curriculum',
    '4. Faculty & Leadership',
    '5. Class Size & Student Profile',
    '6. Admissions Requirements & Your Fit',
    '7. Rankings & Reputation',
    '8. Key Differentiators',
    '9. Notable Kelley Alumni',
    '10. Interviewer Profile: Allison O\'Boyle',
    '11. Interview Questions & Your Personalized Answers',
    '12. Questions YOU Should Ask',
    '13. How to Position DU PhD vs. Kelley DBA',
    '14. How to Present Your Research',
    '15. Red Flags to Avoid',
    '16. Your Key Advantages',
    'Quick Reference Card',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================================
# 1. PROGRAM OVERVIEW & HISTORY
# ============================================================
add_heading_styled('1. Program Overview & History', level=1)

doc.add_paragraph(
    'The Kelley Executive Doctor of Business Administration (EDBA) is a brand-new program. '
    'It was officially announced in January 2026 and will enroll its inaugural cohort in Fall 2026. '
    'There are no prior graduates or existing alumni — you would be in the first class.'
)

add_heading_styled('Why Kelley Launched the EDBA', level=2)
doc.add_paragraph(
    'Kelley identified a gap between the MBA and advanced doctoral-level leadership development for senior executives. '
    'With a 137,000+ alumni network, Kelley received demand signals from senior leaders facing complex challenges: '
    'institutional changes, heightened stakeholder expectations, technological disruption, and geopolitical pressures.'
)

p = doc.add_paragraph()
run = p.add_run('Dean Pat Hopkins stated: ')
run.bold = True
p.add_run('"[The program] builds on this tradition by elevating strategic leadership development through '
          'the integration of advanced business theory and practical application."')

doc.add_paragraph(
    'The program is delivered through Kelley Direct Programs, which also runs the #1 ranked online MBA program '
    '(U.S. News, five consecutive years), giving the EDBA immediate infrastructure credibility.'
)

# ============================================================
# 2. PROGRAM STRUCTURE & FORMAT
# ============================================================
doc.add_page_break()
add_heading_styled('2. Program Structure & Format', level=1)

add_table(
    ['Detail', 'Information'],
    [
        ['Duration', '2 years (with MBA transfer credits) or 3 years (without prior MBA)'],
        ['Total Credits', '69 credit hours (no MBA) or 42 credit hours (with MBA, up to 27 transferable)'],
        ['Modality', 'Hybrid \u2014 primarily online with required in-person residencies'],
        ['Live Sessions', 'One night per week, ~75 minutes, 6:00\u20139:00 PM Eastern Time (recorded)'],
        ['Weekly Commitment', 'Approximately 15\u201320 hours per week'],
        ['Tuition', '$2,300 per credit hour'],
        ['Total Cost (with MBA)', '~$96,600 plus residency fees'],
        ['Total Cost (without MBA)', '~$142,440 (MBA credits at MBA rate, remainder at EDBA rate)'],
        ['Residency Fees', '$1,000 per residency (additional)'],
        ['Concentration', 'Strategy and Leadership'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('YOUR SITUATION: ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('You hold an MBA (Finance) from the University of Denver. This qualifies you for the 2-year, '
          '42-credit-hour track at approximately $96,600 plus residency fees. Up to 27 MBA credits may transfer '
          '(subject to transcript review).')

add_heading_styled('Residency Details', level=2)

doc.add_paragraph('For students WITH a prior MBA (your track) \u2014 2 doctoral residencies:', style='List Bullet')
doc.add_paragraph(
    'EDBA Special Topics I (3 credits) \u2014 Executive presence, persuasion, data-driven communication. '
    'First cohort dates: November 5\u20139, 2026',
    style='List Bullet 2'
)
doc.add_paragraph(
    'EDBA Special Topics II (3 credits) \u2014 Thesis development: refining research questions, methods, '
    'identifying faculty advisors. Summer between year one and two.',
    style='List Bullet 2'
)

# ============================================================
# 3. CURRICULUM
# ============================================================
doc.add_page_break()
add_heading_styled('3. Curriculum', level=1)

doc.add_paragraph('Your 42 doctoral credit hours break into four sections:')

add_heading_styled('A. Strategy & Leadership Courses \u2014 18 Credit Hours', level=2)
doc.add_paragraph(
    'Six courses focused on strategy and leadership. Faculty expertise includes: global strategy, '
    'corporate innovation, leadership, entrepreneurship, international business, consulting, negotiations, '
    'change management, organizational design, crisis management, and turnaround management.'
)

add_heading_styled('B. Statistics & Research Methods \u2014 9 Credit Hours', level=2)
doc.add_paragraph(
    'Two statistics and research methods courses plus one thesis development course. '
    'Covers both quantitative and qualitative research methodologies.'
)

add_heading_styled('C. Thesis / Applied Research Project \u2014 9 Credit Hours', level=2)
doc.add_paragraph(
    'Required applied research thesis with one-on-one faculty advising. Full research lifecycle: '
    'ideation, problem framing, data collection, analysis, writing, and committee presentation. '
    'Students have significant latitude in choosing topics \u2014 thesis should be tailored to the '
    'student\'s own organization or career context.'
)

add_heading_styled('D. EDBA Residencies \u2014 6 Credit Hours', level=2)
doc.add_paragraph('Two in-person residencies at 3 credits each (described above).')

# ============================================================
# 4. FACULTY & LEADERSHIP
# ============================================================
doc.add_page_break()
add_heading_styled('4. Faculty & Leadership', level=1)

doc.add_paragraph(
    'The EDBA is taught by the same full-time, award-winning faculty who teach in the '
    'Kelley Direct MBA program (the #1 online MBA).'
)

add_heading_styled('Erik Gonzalez-Mule, Ph.D. \u2014 Associate Chair, Kelley Direct Programs', level=2)
doc.add_paragraph(
    'Holds the Randall L. Tobias Chair in Leadership. Chairperson, Department of Management and '
    'Entrepreneurship. Ph.D. in Business Administration, University of Iowa (Presidential Research Fellow). '
    'Research: team composition and processes, stress, individual differences, employee health. '
    'Published in Journal of Applied Psychology, Journal of Management. '
    'Named 2024 Poets&Quants Best 40-Under-40 MBA Professor. '
    'Research featured in The Wall Street Journal, CBS News, Forbes.'
)

add_heading_styled('Patrick E. Hopkins, Ph.D. \u2014 Dean, Kelley School of Business', level=2)
doc.add_paragraph(
    'Effective March 17, 2025. Kelley faculty since 1995; previously Vice Dean. '
    'Ph.D. in Accounting, University of Texas at Austin (minors in psychology and statistics). '
    'CPA; won the Elijah Watt Sells Award (1987). Research: professional judgment, decision-making.'
)

add_heading_styled('Other Senior Leadership', level=2)
doc.add_paragraph('Dan Li \u2014 Executive Associate Dean; L. Leslie Waters Chair in International Business')
doc.add_paragraph('Stephanie Lu Wang \u2014 Samuel & Pauline Glaubinger Professor; Associate Chair, Mgmt & Entrepreneurship')
doc.add_paragraph('Will Geoghegan \u2014 Kelley Direct Chair, Clinical Professor')
doc.add_paragraph('Miki Pike Hamstra \u2014 Executive Director of Kelley Direct')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Note: ')
run.bold = True
p.add_run('The Management and Entrepreneurship Department alone has 50+ faculty members.')

# ============================================================
# 5. CLASS SIZE & STUDENT PROFILE
# ============================================================
doc.add_page_break()
add_heading_styled('5. Class Size & Student Profile', level=1)

add_table(
    ['Detail', 'Information'],
    [
        ['Target first cohort', '~25 students'],
        ['Maximum per admissions cycle', 'Capped at 50'],
        ['Target experience level', '10+ years of management experience'],
        ['Target profile', 'Mid-to-late career professionals; senior executives'],
        ['Geographic focus', 'Primarily U.S.-based (evening ET classes), international welcome'],
        ['Industries', 'Diverse (program emphasizes cross-industry peer networking)'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('YOUR FIT: ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('You exceed the target profile significantly: 25+ years of executive experience (vs. 10+ required), '
          'spanning government (Governor\'s office, U.S. Congress), private sector (tech startups, consulting), '
          'and healthcare technology. Your cross-sector breadth is exactly what enriches a cohort.')

# ============================================================
# 6. ADMISSIONS REQUIREMENTS & YOUR FIT
# ============================================================
add_heading_styled('6. Admissions Requirements & Your Fit', level=1)

add_table(
    ['Requirement', 'Details', 'Your Status'],
    [
        ['GMAT/GRE', 'Optional', 'Not needed'],
        ['Work Experience', '10+ years management', '25+ years \u2014 far exceeds'],
        ['Prior Degree', 'MBA preferred', 'MBA (Finance), University of Denver'],
        ['Additional Credentials', '\u2014', 'BBA (Ole Miss) + LSE Certificate'],
        ['Application Fee', 'Phase 1 is FREE', '\u2014'],
        ['Priority Deadline', 'May 1, 2026', 'Well within window'],
    ]
)

add_heading_styled('Application Process', level=2)
doc.add_paragraph(
    'Phase 1: ~5 minutes; answer a few questions and upload resume (free). '
    'If a strong fit, contacted within 2 weeks. '
    'Phase 2: 30-minute interview with admissions team member. '
    'The FAQ states: "No need to be nervous \u2014 we love getting to know our applicants!"'
)

# ============================================================
# 7. RANKINGS & REPUTATION
# ============================================================
doc.add_page_break()
add_heading_styled('7. Rankings & Reputation', level=1)

add_table(
    ['Category', 'Ranking'],
    [
        ['Online MBA', '#1 \u2014 U.S. News, 5 consecutive years (same Kelley Direct platform as EDBA)'],
        ['Full-Time MBA', '#22 (tie) \u2014 U.S. News 2026'],
        ['Part-Time MBA', '#16 \u2014 U.S. News 2026'],
        ['Undergraduate Business', '#8 overall, #4 among publics \u2014 U.S. News 2026'],
        ['Research (UT-Dallas)', '#10 worldwide, #9 North America, #2 among publics'],
        ['Management Research (TAMUGA)', '#4 in U.S., #1 in Big Ten (tied with Michigan)'],
    ]
)

doc.add_paragraph()
add_key_point(
    'There are no established DBA-specific rankings. The program\'s credibility rests on '
    'AACSB accreditation, the #1 online MBA platform, and top-10 worldwide research productivity.'
)

# ============================================================
# 8. KEY DIFFERENTIATORS
# ============================================================
add_heading_styled('8. Key Differentiators', level=1)

differentiators = [
    ('Built on #1 Online MBA Platform', 'Same Kelley Direct infrastructure \u2014 proven technology, pedagogy, and support.'),
    ('Strategy & Leadership Focus', 'Single concentration provides depth over breadth. Not a broad survey.'),
    ('Scholar-Practitioner Model', 'Trains executives to be "critical consumers of information and data" who "question and improve organizations."'),
    ('Applied Thesis (Not Just Capstone)', 'Full applied research thesis \u2014 9 dedicated credit hours with one-on-one faculty advising.'),
    ('Top-10 Research Faculty', 'UT-Dallas top 10 worldwide research. Students learn from actively publishing scholars.'),
    ('Lifetime Career Coaching', 'Kelley Direct Professional Advancement (KDPA) \u2014 lifetime access to one-on-one coaching.'),
    ('Flexible for Executives', 'One 75-min live class/week, recorded sessions, 15-20 hrs/week, 2-4 days residency/year.'),
    ('Cost Competitive', '~$96K\u2013$142K. Competitive with Temple Fox, less than many European programs.'),
    ('Massive Alumni Network', '137,000+ Kelley alumni across 105 countries.'),
    ('AACSB Accredited', 'Full AACSB accreditation at the institutional level.'),
]

for t, desc in differentiators:
    p = doc.add_paragraph()
    run = p.add_run(f'{t}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 9. NOTABLE KELLEY ALUMNI
# ============================================================
doc.add_page_break()
add_heading_styled('9. Notable Kelley Alumni', level=1)

p = doc.add_paragraph()
run = p.add_run('Important: The EDBA has NO alumni yet \u2014 the first cohort starts Fall 2026. ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('However, the broader Kelley network includes:')

alumni = [
    ('Mark Cuban', 'Billionaire entrepreneur, Shark Tank (BS Management \'81)'),
    ('David Solomon', 'CEO, Goldman Sachs (MBA \'92)'),
    ('John T. Chambers', 'Former CEO & Executive Chairman, Cisco Systems (MBA \'75)'),
    ('Alicia Boler Davis', 'CEO, Alto Pharmacy; former GM executive (MBA)'),
    ('Derica Rice', 'Former EVP, CVS Health; Board at Disney, Target, BMS (MBA \'90)'),
    ('Anton Vincent', 'President, Mars Wrigley North America (MBA \'93)'),
    ('Patrick Decker', 'President & CEO, Xylem'),
]
for name, desc in alumni:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{name} \u2014 ')
    run.bold = True
    p.add_run(desc)

# ============================================================
# 10. INTERVIEWER PROFILE
# ============================================================
doc.add_page_break()
add_heading_styled('10. Interviewer Profile: Allison O\'Boyle', level=1)

p = doc.add_paragraph()
run = p.add_run('NOTE: ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('The interviewer\'s name is likely "Allison O\'Boyle" (not "Boyle"). Her IU email is asoboyle@iu.edu.')

doc.add_paragraph()
add_table(
    ['Detail', 'Information'],
    [
        ['Full Name', 'Allison O\'Boyle'],
        ['Title', 'EDBA Admissions and Academic Advisor'],
        ['Email', 'asoboyle@iu.edu'],
        ['Department', 'Kelley Direct Programs'],
        ['Location', '1275 E. Tenth Street, Suite 2100, Bloomington, IN 47405'],
        ['Role in Interview', 'She is the ONLY person listed specifically for EDBA admissions'],
    ]
)

add_heading_styled('What We Know', level=2)
doc.add_paragraph(
    'Allison holds a dual role as both EDBA Admissions contact and Academic Advisor. '
    'She is the dedicated point person for the EDBA program and the only person listed specifically '
    'for EDBA admissions on the Kelley Direct contact page.'
)

add_heading_styled('Reports To / Works With', level=2)
doc.add_paragraph('Abigail Gschwend \u2014 Director of Admissions and Financial Aid (likely her boss)', style='List Bullet')
doc.add_paragraph('Courtney Daily \u2014 Associate Director of Enrollment Management', style='List Bullet')
doc.add_paragraph('Erik Gonzalez-Mule \u2014 Associate Chair, Kelley Direct (EDBA academic lead)', style='List Bullet')
doc.add_paragraph('Miki Pike Hamstra \u2014 Executive Director of Kelley Direct', style='List Bullet')

add_heading_styled('Possible Family Connection', level=2)
doc.add_paragraph(
    'There is a notable faculty member at Kelley named Ernest O\'Boyle (eoboyle@iu.edu), '
    'who is the Dale M. Coleman Chair of Management and Professor in the Department of Management '
    'and Entrepreneurship. Ph.D. from Virginia Commonwealth University. Research: individual differences, '
    'Dark Triad personality, research methods. The shared last name could indicate a family relationship, '
    'though this is unconfirmed. If related, she likely has deep familiarity with academic research culture.'
)

add_heading_styled('What to Expect from Her', level=2)
doc.add_paragraph('Conversational, not adversarial \u2014 the FAQ literally says "we love getting to know our applicants"', style='List Bullet')
doc.add_paragraph('Fit-focused \u2014 she wants to understand your trajectory and why this program, why now', style='List Bullet')
doc.add_paragraph('Program-selling too \u2014 as recruiter for a brand-new program, she wants strong candidates', style='List Bullet')
doc.add_paragraph('Research-aware \u2014 as academic advisor for a doctoral program, she appreciates hearing research interests', style='List Bullet')
doc.add_paragraph('30 minutes \u2014 be concise and prepared with your narratives', style='List Bullet')

# ============================================================
# 11. INTERVIEW QUESTIONS & YOUR PERSONALIZED ANSWERS
# ============================================================
doc.add_page_break()
add_heading_styled('11. Interview Questions & Your Personalized Answers', level=1)

p = doc.add_paragraph()
run = p.add_run('These answers are crafted from your actual resume and career history. '
                'Adapt the language to feel natural, but the substance is drawn directly from your experience.')
run.italic = True

# --- MOTIVATION ---
doc.add_page_break()
add_heading_styled('A. Motivation & "Why" Questions', level=2)

add_qa(
    '"Tell me about yourself and why you\'re interested in the Executive DBA."',
    'I\'ve spent over 25 years at the intersection of business strategy, government relations, and '
    'technology. I started in the Mississippi Governor\'s office, where I was part of the team that secured '
    'the $1.6 billion Nissan manufacturing plant \u2014 the largest economic development project in the state\'s '
    'history, creating 30,000 jobs. From there I went to Congress, where I served as Deputy Chief of Staff '
    'and co-authored sections of the Dodd-Frank Act on the House Financial Services Committee. Then I moved '
    'into the private sector \u2014 spending nearly a decade at GetInsured leading government technology RFP '
    'responses that generated over $150 million in revenue, then helping Sidecar Health scale to a $1 billion '
    'valuation. Today I run my own strategic consulting firm where I\'ve driven over $200 million in revenue '
    'through RFP strategy and government relations.\n\n'
    'Through all of that, one question has followed me: does the way government evaluates proposals actually '
    'lead to better outcomes? I\'ve seen lowest-price contracting produce terrible results and best-value '
    'approaches deliver real public impact. But there\'s almost no rigorous empirical research testing this. '
    'I want to be the person who produces that research \u2014 and the Executive DBA\'s focus on applied, '
    'practitioner-driven scholarship is the perfect vehicle for that.'
)

add_qa(
    '"Why now? What makes this the right time in your career?"',
    'Three things converge right now. First, I\'ve accumulated 25+ years of direct experience in the '
    'exact domain I want to study \u2014 government procurement, RFP strategy, legislative influence on '
    'contracting outcomes. I\'ve lived this from the policy side in the Governor\'s office and Congress, '
    'from the vendor side at GetInsured and Sidecar Health, and from the advisory side at my own firm. '
    'That breadth gives me a research perspective that\'s rare.\n\n'
    'Second, I\'ve already begun the research. I\'ve built data pipelines from USAspending.gov and FPDS, '
    'run preliminary statistical analyses, and developed my research framework. I\'m not starting from '
    'zero \u2014 I\'m bringing a working research program.\n\n'
    'Third, the Kelley EDBA launching its inaugural cohort right now is genuinely exciting. I\'ve built '
    'things from the ground up my entire career \u2014 from startups achieving 150-700% growth to '
    'transforming Mississippi\'s entire economic development apparatus. Being part of a first cohort is '
    'exactly the kind of opportunity I gravitate toward.'
)

add_qa(
    '"Why a DBA rather than a PhD?"',
    'My goal isn\'t to become a full-time academic \u2014 it\'s to produce research that changes how '
    'practitioners operate. The DBA\'s applied thesis model, where I can use real federal contract data '
    'to answer a question that directly affects how hundreds of billions in taxpayer dollars are allocated, '
    'fits that goal precisely. I\'ve spent my career as a practitioner who uses evidence to drive decisions \u2014 '
    'from transforming raw data into Tableau dashboards that boosted sales 25% in a month, to building '
    'cost-tracking systems that replaced estimate-based pricing in RFPs. The scholar-practitioner identity '
    'isn\'t aspirational for me \u2014 it\'s who I already am. The DBA gives me the doctoral framework to '
    'formalize that.'
)

add_qa(
    '"Why Kelley specifically?"',
    'Four reasons. First, the #1-ranked online MBA platform tells me the delivery infrastructure is '
    'proven \u2014 I\'m not a test case for untested technology. Second, the strategy and leadership '
    'concentration maps directly to my research: procurement source selection is fundamentally a strategic '
    'organizational decision. Third, Kelley\'s top-10 worldwide research ranking means I\'ll be learning '
    'from faculty who are actively publishing at the highest levels. And fourth, the opportunity to be in '
    'the inaugural cohort \u2014 to help define what this program becomes \u2014 is compelling. I\'ve spent my '
    'career building things. The idea of helping shape an executive doctoral program at a top-10 research '
    'institution is exactly the kind of challenge I seek out.'
)

add_qa(
    '"What do you hope to achieve that you cannot without this degree?"',
    'Credibility and rigor. I can write a white paper on procurement reform tomorrow based on my '
    'experience. But to produce research that changes federal acquisition policy, I need doctoral-level '
    'methodology, peer review, and the institutional weight of a top research university behind me. '
    'I want my findings published in journals that policymakers and agency leaders actually read. '
    'The DBA gives me the methodological training, faculty mentorship, and academic credibility to make '
    'that happen. Without it, I\'m just another consultant with opinions. With it, I\'m a scholar-practitioner '
    'with evidence.'
)

# --- PROFESSIONAL BACKGROUND ---
doc.add_page_break()
add_heading_styled('B. Professional Background Questions', level=2)

add_qa(
    '"Walk us through your career and how it led you here."',
    'I\'ll give you the arc. I started in state government \u2014 Deputy Chief of Staff in the Mississippi '
    'Governor\'s office from 2000 to 2004. That\'s where I got my first taste of how government decisions '
    'drive economic outcomes. I was on the team that landed the $1.6 billion Nissan plant in Canton \u2014 '
    '6,000 direct jobs, 24,000 ancillary jobs, the largest economic development project in state history. '
    'I also managed over $700 million annually in Community Development Block Grants and rewrote the state\'s '
    'economic development policy, which led to creating the Mississippi Development Authority.\n\n'
    'After running my own public policy consulting firm, I went to Washington as Deputy Chief of Staff for '
    'Congressman Travis Childers on the House Financial Services Committee. I co-authored sections of the '
    'Dodd-Frank Act \u2014 specifically the credit card and loan document provisions. That experience taught me '
    'how legislation shapes markets.\n\n'
    'Then I moved to the private sector. At GetInsured, I spent nearly a decade leading government technology '
    'RFP responses \u2014 securing over $150 million in new revenue. At Sidecar Health, I transformed their '
    'distribution model and helped drive them to a $1 billion valuation. Today I run Davidson Strategic '
    'Services, where I\'ve generated over $200 million in revenue through RFP strategy and government relations, '
    'including $65 million in state and federal contracts.\n\n'
    'The through-line is this: I\'ve always operated at the intersection of government decision-making and '
    'business outcomes. The DBA lets me study that intersection rigorously.'
)

add_qa(
    '"What is the most significant leadership challenge you have faced?"',
    'At Sidecar Health, I was brought in to transform the entire go-to-market strategy. They were a '
    'healthcare technology startup relying solely on SEO-driven direct sales. I had to build a nationwide '
    'partner distribution network from scratch \u2014 design the framework, recruit and onboard thousands of '
    'distribution partners, create training and certification programs, and build a sales enablement team.\n\n'
    'The biggest challenge was speed. We developed a self-service registration portal that reduced partner '
    'onboarding from 21+ days to under 5 hours, which increased our partner pipeline by 1,300% in two months. '
    'I also launched an e-learning platform that drove a 30% sales increase in established markets. The result '
    'was explosive growth that contributed directly to their $1 billion valuation and successful funding round.\n\n'
    'The leadership lesson: you have to build systems that scale, not just make individual deals. That same '
    'insight drives my research interest \u2014 procurement evaluation systems determine outcomes at massive scale.'
)

add_qa(
    '"Describe a time you used data or evidence to make an important decision."',
    'At Davidson Strategic Services, a client was struggling with flat sales despite a strong product. '
    'I took their raw sales data and built Tableau dashboards that revealed patterns they couldn\'t see \u2014 '
    'geographic clusters of underperformance, timing patterns in the sales cycle, and pricing sensitivity '
    'by segment. Within one month, those insights boosted existing sales by 25%.\n\n'
    'But the bigger example is what I\'m doing now with my research. I\'ve built data pipelines pulling '
    'federal contract data from USAspending.gov, FPDS records, and GAO bid protest data. I\'ve run '
    'OLS regression, negative binomial models, logistic regression, and propensity score matching on '
    '15,000+ contract records. The preliminary results show meaningful differences between lowest-price '
    'and best-value procurement outcomes. That\'s the kind of evidence-based analysis I want to formalize '
    'through the DBA.'
)

# --- RESEARCH ---
doc.add_page_break()
add_heading_styled('C. Research Interest Questions', level=2)

add_qa(
    '"What are your research interests?"',
    'My dissertation topic is "From Lowest Price to Highest Public Value: An Empirical Test of Best-Value '
    'Source Selection in Government RFPs." The U.S. federal government spends over $700 billion annually '
    'in contracts. The default approach \u2014 awarding to the lowest bidder \u2014 often produces poor outcomes: '
    'cost overruns, bid protests, and failed contractor performance. Best-value source selection evaluates '
    'both price and non-price factors like technical capability and past performance, but we lack rigorous '
    'empirical evidence on whether it actually delivers better results.\n\n'
    'I\'m filling that gap. I\'ve already collected FY2024 data from USAspending.gov (15,477 analysis records), '
    'built the statistical models, and run preliminary analyses. This isn\'t hypothetical \u2014 it\'s underway.\n\n'
    'What makes this personal is that I\'ve lived on both sides. I\'ve written 19+ successful government '
    'technology RFPs that generated $200M+ in revenue. I\'ve navigated the evaluation process as a vendor. '
    'And I\'ve influenced procurement policy from inside government. I know this system from every angle.'
)

add_qa(
    '"How would your research create value for practitioners?"',
    'If my research shows that best-value source selection produces measurably better outcomes \u2014 fewer '
    'protests, better contractor performance, lower cost overruns \u2014 then procurement leaders have empirical '
    'justification to use it more aggressively, even when it\'s administratively harder. If it doesn\'t, we '
    'need to understand why the theory fails in practice.\n\n'
    'Either way, the findings directly change how contracting officers and agency leaders make decisions '
    'about hundreds of billions in taxpayer dollars. And the underlying question \u2014 does paying for quality '
    'produce better outcomes than minimizing upfront cost? \u2014 applies to any organization making purchasing '
    'decisions. It\'s a universal strategy question with a very specific, testable government context.'
)

add_qa(
    '"What data or resources do you have for your thesis?"',
    'I\'ve already built a substantial research infrastructure. I pull contract data from the USAspending.gov '
    'API (no authentication needed, publicly available), FPDS contract records for detailed procurement '
    'methodology data, and GAO bid protest records. I\'ve processed FY2024 Q2 data \u2014 15,477 analysis-ready '
    'records. My statistical toolkit includes Python with Jupyter notebooks, running OLS regression, negative '
    'binomial generalized linear models, logistic regression, and propensity score matching. I also have '
    'Tableau experience for visualization. I\'m not starting from scratch \u2014 I\'m bringing a working research '
    'program to the program.'
)

# --- CAPACITY & COMMITMENT ---
doc.add_page_break()
add_heading_styled('D. Capacity & Commitment Questions', level=2)

add_qa(
    '"How do you plan to balance the program with professional responsibilities?"',
    'I run my own consulting firm, Davidson Strategic Services, which gives me significant control over '
    'my schedule. I\'ve intentionally structured my business to allow for this kind of commitment. The '
    'hybrid format \u2014 one 75-minute live session per week in the evenings, plus 15-20 hours of coursework \u2014 '
    'fits cleanly into my working life.\n\n'
    'More importantly, I\'ve managed much heavier operational loads before. When I was at GetInsured, I was '
    'simultaneously managing government RFP responses, stakeholder relationships, and legislative initiatives '
    'across multiple states. At Sidecar Health, I built an entire nationwide distribution network while '
    'managing onboarding for thousands of partners. I\'m not someone who struggles with volume \u2014 I organize '
    'and execute.'
)

add_qa(
    '"The program requires significant time commitment. Are you prepared for that?"',
    'Absolutely. I want to be direct: I\'m not entering this casually. I\'ve already invested hundreds of '
    'hours in my research \u2014 building data pipelines, running preliminary analyses, writing a research '
    'framework. I chose to pursue a doctorate because I have a specific question I need answered, not because '
    'I\'m looking for a credential. That intrinsic motivation is what sustains commitment when the workload '
    'is heavy. I\'ve also operated in high-intensity environments my entire career \u2014 from managing Congressional '
    'operations during the 2008 financial crisis to driving 150-700% year-over-year growth at startups. '
    'I know what sustained high performance looks like.'
)

# --- KELLEY-SPECIFIC ---
doc.add_page_break()
add_heading_styled('E. Kelley-Specific & Fit Questions', level=2)

add_qa(
    '"What does being a scholar-practitioner mean to you?"',
    'It means using rigorous evidence to make better decisions \u2014 and producing research that practitioners '
    'can actually use. I\'ve been doing this informally my whole career. When I built Tableau dashboards '
    'that increased sales 25% in a month, that was applied analytics. When I replaced cost estimates in '
    'RFPs with a detailed cost-tracking system, that was evidence-based process improvement. When I '
    'authored sections of the Dodd-Frank Act, I was translating policy research into legislation.\n\n'
    'The difference the DBA makes is formalization. Instead of relying on experience and intuition backed '
    'by ad hoc data, I want to apply proper research methodology \u2014 rigorous statistical analysis, peer '
    'review, replicable methods \u2014 to questions that affect billions in public spending. A scholar-practitioner '
    'doesn\'t just lead; they generate the knowledge that makes leadership evidence-based.'
)

add_qa(
    '"How does strategy and leadership relate to your research?"',
    'Procurement source selection is one of the highest-stakes strategic decisions a government agency makes. '
    'When a contracting officer chooses between lowest-price technically acceptable and best-value tradeoff, '
    'they\'re making a strategic resource allocation decision that determines whether the government \u2014 and '
    'by extension, taxpayers \u2014 gets the best possible outcome.\n\n'
    'I\'ve seen this from the leadership side too. At the Governor\'s office, I watched how the strategic '
    'decision to pursue the Nissan plant required a complete overhaul of Mississippi\'s economic development '
    'infrastructure. In Congress, I saw how legislative strategy on the Financial Services Committee shaped '
    'market outcomes for decades. Strategy and leadership aren\'t abstract concepts for me \u2014 they\'re the '
    'operational reality of every role I\'ve held.'
)

add_qa(
    '"What would you contribute to the cohort?"',
    'I bring a perspective that\'s unusual for a business school cohort: deep experience in government at '
    'the state and federal level, combined with private sector technology startup experience. Most of my '
    'peers will come from corporate backgrounds. I bring the public sector lens \u2014 how government makes '
    'decisions, how legislation and regulation shape markets, how public-private partnerships work in practice.\n\n'
    'Specifically: I\'ve co-authored federal legislation (Dodd-Frank), I\'ve secured a $1.6 billion '
    'manufacturing plant for a state, I\'ve built distribution networks that drove a billion-dollar valuation, '
    'and I\'ve generated $200M+ through government RFP strategy. That range of experience \u2014 from the Governor\'s '
    'mansion to Capitol Hill to Silicon Valley to my own consulting firm \u2014 gives me a unique vantage point '
    'that I think enriches the learning environment for everyone.'
)

add_qa(
    '"What excites you about being part of the inaugural cohort?"',
    'I\'ve spent my career building things from scratch. I transformed Mississippi\'s economic development '
    'apparatus into the Mississippi Development Authority. I built Sidecar Health\'s national distribution '
    'network from zero. I built my own consulting practice. The inaugural cohort of the Kelley EDBA is '
    'another chance to build something from the ground up \u2014 but this time in an academic context.\n\n'
    'There\'s also a practical advantage: the inaugural cohort will set the tone and culture for every '
    'cohort that follows. That\'s a meaningful opportunity to shape something lasting. And frankly, Kelley '
    'is investing its best resources in making this first cohort a success. Being part of that concentrated '
    'investment in quality is appealing.'
)

add_qa(
    '"What are your strengths and weaknesses?"',
    'My strength is execution. I don\'t just strategize \u2014 I build. Whether it\'s writing the RFP response, '
    'standing up the distribution network, or drafting the legislation, I do the work. That\'s reflected in '
    'measurable outcomes: $200M+ in revenue, $1.6B in economic development, a $1B startup valuation.\n\n'
    'My weakness is that I can be impatient with process when I see the answer clearly. In government, '
    'I learned that stakeholder alignment takes time even when the right course of action is obvious. '
    'In the doctoral context, I expect the same lesson applies \u2014 rigorous methodology requires patience '
    'and thoroughness that can feel slow when you\'re eager to reach the conclusion. I\'m aware of that '
    'tendency and actively manage it.'
)

# --- HANDLING THE DU QUESTION ---
doc.add_page_break()
add_heading_styled('F. The DU PhD Question (If It Comes Up)', level=2)

p = doc.add_paragraph()
run = p.add_run('CONTEXT: ')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
p.add_run('You hold an MBA from the University of Denver and have been preparing for a PhD at DU\'s '
          'Daniels College of Business starting Fall 2026. If asked, here is how to frame it:')

add_qa(
    '"Are you considering other programs?"',
    'Yes, and I want to be transparent about that. I\'ve been accepted to a PhD program at the University '
    'of Denver\'s Daniels College of Business \u2014 where I earned my MBA \u2014 also starting Fall 2026. I\'ve been '
    'doing extensive research preparation, including building data pipelines and running preliminary analyses.\n\n'
    'When I learned about the Kelley EDBA, I was drawn to it for specific reasons: the strategy and '
    'leadership concentration maps directly to my research, the applied thesis model fits my practitioner '
    'orientation, and the caliber of Kelley\'s research faculty is extraordinary \u2014 top 10 worldwide. I\'m '
    'taking this decision seriously because I want to choose the program that best serves my research goals '
    'and maximizes the impact of my work on procurement policy.\n\n'
    'Both programs would allow me to pursue my research, but the approaches differ. The PhD is more '
    'traditionally academic; the DBA emphasizes applied impact. I\'m evaluating which approach produces '
    'research that most effectively changes how government agencies award contracts. That\'s why I wanted '
    'to have this conversation.'
)

# ============================================================
# 12. QUESTIONS YOU SHOULD ASK
# ============================================================
doc.add_page_break()
add_heading_styled('12. Questions YOU Should Ask', level=1)

doc.add_paragraph(
    'Pick 3\u20134 that feel most natural. These demonstrate research, curiosity, and strategic thinking.'
)

add_heading_styled('About Research & Thesis', level=2)
doc.add_paragraph(
    '"My research uses federal contract data from USAspending.gov and FPDS to study procurement evaluation '
    'methods. How does the program support students whose research sits at the intersection of public '
    'policy and organizational strategy?"',
    style='List Bullet'
)
doc.add_paragraph(
    '"How are students matched with faculty advisors, and how early in the program does that process begin?"',
    style='List Bullet'
)
doc.add_paragraph(
    '"What does the thesis defense process look like? Is it structured like a traditional dissertation '
    'defense, or does it have a more practitioner-oriented format?"',
    style='List Bullet'
)

add_heading_styled('About the Cohort', level=2)
doc.add_paragraph(
    '"With the inaugural cohort around 25 students, what professional backgrounds and industries are you '
    'seeing in the applicant pool? I\'m curious about the diversity of perspectives."',
    style='List Bullet'
)
doc.add_paragraph(
    '"I\'ve spent my career spanning government and private sector. Is the program seeing interest from '
    'other candidates with public sector or government relations backgrounds?"',
    style='List Bullet'
)

add_heading_styled('About Faculty', level=2)
doc.add_paragraph(
    '"I noticed Professor Gonzalez-Mule\'s research on team composition and organizational behavior. '
    'Are there faculty whose work touches on organizational decision-making, procurement strategy, '
    'or public sector management?"',
    style='List Bullet'
)
doc.add_paragraph(
    '"Will EDBA students have access to research seminars, working papers, or collaboration opportunities '
    'with PhD students?"',
    style='List Bullet'
)

add_heading_styled('About Career Outcomes', level=2)
doc.add_paragraph(
    '"What outcomes is the program hoping to see from graduates? I\'m interested in both applied leadership '
    'impact and the potential to teach or publish."',
    style='List Bullet'
)

add_heading_styled('One Bold Question', level=2)
doc.add_paragraph(
    '"As the inaugural cohort, we\'ll be setting the standard for this program. What\'s your vision for '
    'where the Kelley EDBA will be positioned in five years relative to established executive doctoral programs?"',
    style='List Bullet'
)

# ============================================================
# 13. HOW TO POSITION DU PhD vs. KELLEY DBA
# ============================================================
doc.add_page_break()
add_heading_styled('13. How to Position DU PhD vs. Kelley DBA', level=1)

p = doc.add_paragraph()
run = p.add_run('This is the most delicate part of the interview. Be honest but strategic.')
run.bold = True
run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)

add_heading_styled('Key Facts to Remember', level=2)
doc.add_paragraph('You earned your MBA at DU \u2014 you have an existing relationship with the university.', style='List Bullet')
doc.add_paragraph('The PhD at DU\'s Daniels College is a traditional academic doctorate.', style='List Bullet')
doc.add_paragraph('The Kelley EDBA is an applied, practitioner-focused executive doctorate.', style='List Bullet')
doc.add_paragraph('Both start Fall 2026 \u2014 this is a genuine, time-sensitive decision.', style='List Bullet')

add_heading_styled('What to Emphasize', level=2)
doc.add_paragraph(
    'The DBA format may be a better fit for your career stage \u2014 applied research and scholar-practitioner model '
    'align with producing research that directly improves government contracting practices.',
    style='List Bullet'
)
doc.add_paragraph(
    'The hybrid/executive format works with your consulting practice.',
    style='List Bullet'
)
doc.add_paragraph(
    'Kelley\'s brand: R1 institution, #1 online MBA, top-10 research worldwide.',
    style='List Bullet'
)
doc.add_paragraph(
    'Strategy and leadership concentration maps directly to procurement as strategic decision-making.',
    style='List Bullet'
)

add_heading_styled('The Strongest Frame', level=2)
doc.add_paragraph(
    '"Both programs would allow me to pursue my research, but the approaches differ. The PhD route is more '
    'traditionally academic, and the DBA route emphasizes applied impact and practitioner relevance. I\'m '
    'genuinely evaluating which approach will produce research that most effectively changes how government '
    'agencies award contracts \u2014 and which learning environment will push me the hardest. That\'s why I '
    'wanted to have this conversation."'
)

add_heading_styled('What NOT to Say', level=2)
doc.add_paragraph('Never disparage DU or Daniels \u2014 you earned your MBA there.', style='List Bullet')
doc.add_paragraph('Don\'t say EDBA is a "backup" or "easier path."', style='List Bullet')
doc.add_paragraph('Don\'t frame it as "I got in somewhere else, but I\'d rather go here."', style='List Bullet')
doc.add_paragraph('Don\'t say "I\'m just exploring" \u2014 sounds uncommitted.', style='List Bullet')

# ============================================================
# 14. HOW TO PRESENT YOUR RESEARCH
# ============================================================
doc.add_page_break()
add_heading_styled('14. How to Present Your Research', level=1)

add_heading_styled('The 60-Second Version', level=2)
doc.add_paragraph(
    '"The U.S. federal government spends over $700 billion annually in contracts. The default approach \u2014 '
    'awarding to the lowest bidder \u2014 often produces poor outcomes: cost overruns, protests, and failed '
    'performance. Best-value source selection evaluates both price and non-price factors, but we lack '
    'rigorous empirical evidence on whether it actually delivers better results. My research fills that gap. '
    'I\'ve already collected 15,000+ federal contract records and run preliminary statistical analyses. '
    'This isn\'t hypothetical \u2014 it\'s underway."'
)

add_heading_styled('Why It\'s Personal (Use If Asked)', level=2)
doc.add_paragraph(
    '"I\'ve written 19+ successful government technology RFPs generating over $200 million in revenue. '
    'I\'ve navigated the evaluation process as a vendor. I\'ve influenced procurement policy from inside '
    'government \u2014 in the Governor\'s office, in Congress, and through legislative advocacy. I know this '
    'system from every angle, and I know it can work better."'
)

add_heading_styled('Why It Fits Kelley\'s Strategy & Leadership Focus', level=2)
doc.add_paragraph(
    '"Procurement source selection is fundamentally a strategic decision. When a contracting officer '
    'chooses between lowest-price and best-value, they\'re making a resource allocation decision that '
    'determines outcomes for taxpayer dollars. My research applies evidence-based analysis to one of '
    'the largest purchasing decisions any organization makes."'
)

add_heading_styled('The Universal Relevance', level=2)
doc.add_paragraph(
    '"While my specific context is government procurement, the underlying question \u2014 does paying for '
    'quality produce better outcomes than minimizing upfront cost? \u2014 applies to any organization. '
    'It\'s a question about organizational strategy and resource allocation that every executive faces."'
)

# ============================================================
# 15. RED FLAGS TO AVOID
# ============================================================
doc.add_page_break()
add_heading_styled('15. Red Flags to Avoid', level=1)

red_flags = [
    ('Don\'t appear to want the title without the work',
     'Show genuine intellectual curiosity about the research process, not just the credential.'),
    ('Don\'t be vague about research',
     'You have a specific, well-developed question with data already in hand. Use that advantage.'),
    ('Don\'t ask questions answered on the website',
     'Don\'t ask "How many credits?" or "Is it online?" \u2014 these signal you haven\'t done homework.'),
    ('Don\'t underestimate the rigor',
     'Avoid any language suggesting the DBA is easier than a PhD. The workload is substantial.'),
    ('Don\'t be dismissive of theory',
     'Say: "I want to test theory against real-world data to see what actually works."'),
    ('Don\'t badmouth DU or any other program',
     'Frame your pursuit as moving toward Kelley, not running from DU. You earned your MBA there.'),
    ('Don\'t monologue',
     '30 minutes. Be concise. Answer directly, elaborate briefly. Leave room for conversation.'),
    ('Do show personality',
     'The FAQ says "we love getting to know applicants." Be professional but authentic.'),
    ('Don\'t oversell',
     'Your resume speaks for itself. Let the accomplishments land naturally without inflating them.'),
]

for t, desc in red_flags:
    p = doc.add_paragraph()
    run = p.add_run(f'{t}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    p.add_run(desc)

# ============================================================
# 16. YOUR KEY ADVANTAGES
# ============================================================
doc.add_page_break()
add_heading_styled('16. Your Key Advantages Going In', level=1)

doc.add_paragraph(
    'You are not a typical DBA applicant. Here is what sets you apart:'
)

advantages = [
    ('25+ years spanning government and private sector',
     'Governor\'s office, U.S. Congress, Silicon Valley startups, independent consulting. '
     'The program requires 10+ years. You bring 25+.'),
    ('You\'ve already started the research',
     '15,477 contract records, four statistical models, data pipelines built. '
     'Most applicants walk in with vague ideas. You walk in with preliminary results.'),
    ('Elite credentials',
     'MBA from University of Denver (Finance), BBA from Ole Miss, Certificate from the London School of Economics. '
     'Plus the MBA means you qualify for the accelerated 2-year, $96K track.'),
    ('Quantifiable impact at massive scale',
     '$200M+ in RFP revenue, $1.6B Nissan plant (30,000 jobs), $1B startup valuation, '
     'co-author of the Dodd-Frank Act, $65M+ in government contracts.'),
    ('RFP expertise IS your research topic',
     'You\'ve written 19+ successful government technology RFPs. You don\'t just study procurement \u2014 '
     'you\'ve lived it from every angle: policy maker, vendor, and consultant.'),
    ('Public sector perspective diversifies the cohort',
     'Business school cohorts skew private sector. Your Governor\'s office, Congressional, and government '
     'relations experience brings a perspective most classmates won\'t have.'),
    ('You build things from scratch',
     'Mississippi Development Authority, Sidecar Health\'s national distribution network, your own firm. '
     'The inaugural cohort needs builders. That\'s you.'),
    ('Technical research capability',
     'Python, Jupyter notebooks, Tableau, API data pipelines, statistical modeling. '
     'You\'re not just an executive \u2014 you\'re hands-on with data.'),
]

for title, desc in advantages:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(desc)
    doc.add_paragraph()

# ============================================================
# QUICK REFERENCE CARD (LAST PAGE)
# ============================================================
doc.add_page_break()
add_heading_styled('QUICK REFERENCE CARD \u2014 Interview Day', level=1)

info_items = [
    ('Interviewer', 'Allison O\'Boyle, EDBA Admissions & Academic Advisor (asoboyle@iu.edu)'),
    ('Format', '30-minute interview (conversational, not adversarial)'),
    ('Program', 'Kelley Executive DBA \u2014 Strategy & Leadership concentration'),
    ('Your Track', '2 years, 42 credits, ~$96,600 (MBA from DU transfers up to 27 credits)'),
    ('Key Fact', 'Inaugural cohort (Fall 2026) \u2014 you\'d be in the FIRST class'),
    ('Platform', 'Kelley Direct \u2014 same as #1 ranked online MBA (5 consecutive years)'),
    ('Academic Lead', 'Erik Gonzalez-Mule, Tobias Chair in Leadership'),
    ('Dean', 'Pat Hopkins (since March 2025)'),
    ('Priority Deadline', 'May 1, 2026 (rolling admissions)'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(value)

doc.add_paragraph()
add_heading_styled('Your 30-Second Pitch', level=2)
doc.add_paragraph(
    '"I\'ve spent 25 years at the intersection of government decision-making and business outcomes \u2014 from '
    'the Governor\'s office in Mississippi, where I helped land a $1.6 billion Nissan plant, to the U.S. Congress, '
    'where I co-authored sections of the Dodd-Frank Act, to the private sector, where I\'ve generated over '
    '$200 million through government RFP strategy.\n\n'
    'One question has followed me through all of it: does the way government evaluates proposals actually '
    'produce better outcomes? I\'ve started building the empirical research to answer that \u2014 15,000+ federal '
    'contract records, preliminary statistical models, a clear research framework. What I need now is the '
    'doctoral rigor, faculty mentorship, and institutional credibility to turn that research into something '
    'that changes federal procurement policy.\n\n'
    'The Kelley EDBA\'s focus on strategy and leadership, the applied thesis model, and the caliber of your '
    'research faculty make this the right program. And the chance to be part of the inaugural cohort \u2014 to '
    'help build something from the ground up \u2014 is exactly the kind of opportunity I\'ve pursued my entire career."'
)

doc.add_paragraph()
add_heading_styled('Your Key Numbers (Memorize These)', level=2)

numbers = [
    ('$200M+', 'Revenue generated through RFP strategy (19+ successful responses)'),
    ('$1.6B', 'Nissan plant you helped secure for Mississippi (30,000 jobs)'),
    ('$1B', 'Sidecar Health valuation you helped drive'),
    ('$150M+', 'Revenue from government technology RFPs at GetInsured'),
    ('$65M+', 'State and federal contracts through government relations'),
    ('$700M/yr', 'Community Development Block Grants you managed'),
    ('15,477', 'Federal contract records already collected for your research'),
    ('25+', 'Years of professional experience'),
    ('1,300%', 'Partner pipeline increase at Sidecar Health (21+ days to 5 hours onboarding)'),
    ('150\u2013700%', 'Year-over-year growth achieved for startups'),
]

for num, desc in numbers:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  \u2014  ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
    p.add_run(desc)

doc.add_paragraph()
add_heading_styled('Final Reminders', level=2)
doc.add_paragraph('Be yourself \u2014 they want to know YOU, not a performance.', style='List Bullet')
doc.add_paragraph('You are also evaluating THEM \u2014 this is a mutual fit conversation.', style='List Bullet')
doc.add_paragraph('Your preparation puts you ahead of virtually any applicant they\'ll see.', style='List Bullet')
doc.add_paragraph('Ask your questions \u2014 it shows engagement and seriousness.', style='List Bullet')
doc.add_paragraph('Don\'t rush \u2014 30 minutes goes fast. Let the conversation breathe.', style='List Bullet')

# Save
output_path = '/Users/richardwdavidson/Desktop/Kelley_EDBA_Interview_Prep_v2.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
