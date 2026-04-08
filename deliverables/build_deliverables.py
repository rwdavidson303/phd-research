#!/usr/bin/env python3
"""
Build PhD Pivot Deliverables as Word Documents.

Generates 4 Word documents from markdown source files:
  1. Qualitative Pivot Justification Memo (academic: Times New Roman, APA)
  2. Target Journal Strategy (report: Calibri, colored headers)
  3. Case Study Research Design (academic: Times New Roman, APA)
  4. Qualitative Software Guide (report: Calibri, colored headers)
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DELIVERABLES_DIR = Path(__file__).parent
OUTPUT_DIR = DELIVERABLES_DIR / 'output'
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Color constants
DARK_BLUE = RGBColor(0x2E, 0x40, 0x57)
MEDIUM_BLUE = RGBColor(0x1A, 0x5C, 0x8A)
DARK_BLUE_HEX = '2E4057'
ALT_ROW_HEX = 'F0F4F8'


# ============================================================
# Utility functions
# ============================================================

def clean_text(text):
    """Replace HTML entities and markdown artifacts."""
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&mdash;', '\u2014')
    text = text.replace('&ndash;', '\u2013')
    # Convert markdown dashes only in running text, not table separators
    text = re.sub(r'(?<!\|)---(?!\|)', '\u2014', text)
    text = re.sub(r'(?<!\|)--(?!\|)', '\u2013', text)
    return text


def strip_markdown(text):
    """Remove markdown bold/italic/code markers."""
    text = clean_text(text)
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    tc_pr.append(shd)


def apply_inline_formatting(paragraph, text, font_name='Calibri', font_size=Pt(10)):
    """Apply bold, italic, and code formatting to a paragraph from markdown text."""
    text = clean_text(text)
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            continue  # Skip default font assignment
        else:
            run = paragraph.add_run(part)
        run.font.name = font_name
        run.font.size = font_size


def parse_line(line):
    """Classify a markdown line."""
    stripped = line.strip()

    if stripped in ('---', '***', '___'):
        return ('hr', 0, '')
    if stripped.startswith('#'):
        level = 0
        for ch in stripped:
            if ch == '#':
                level += 1
            else:
                break
        return ('heading', level, stripped[level:].strip())
    if not stripped:
        return ('empty', 0, '')
    if stripped.startswith('|') and stripped.endswith('|'):
        return ('table_row', 0, stripped)
    if stripped.startswith('- [ ] ') or stripped.startswith('- [x] '):
        return ('checklist', 0, stripped[2:])
    if stripped.startswith('- ') or stripped.startswith('* '):
        return ('bullet', 0, stripped[2:])
    match = re.match(r'^(\d+)\.\s+(.+)', stripped)
    if match:
        return ('numbered', int(match.group(1)), match.group(2))
    if stripped.startswith('```'):
        return ('code_fence', 0, stripped[3:])
    return ('paragraph', 0, stripped)


# ============================================================
# Academic document (Times New Roman, APA-style)
# ============================================================

def create_academic_doc():
    """Create a Word document with academic formatting."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Configure heading styles (APA)
    for level in range(1, 6):
        sname = f'Heading {level}'
        if sname in doc.styles:
            hs = doc.styles[sname]
            hs.font.name = 'Times New Roman'
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs.font.size = Pt(12)
            hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            if level == 1:
                hs.font.bold = True
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hs.paragraph_format.space_before = Pt(18)
                hs.paragraph_format.space_after = Pt(12)
            elif level == 2:
                hs.font.bold = True
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(6)
            elif level == 3:
                hs.font.bold = True
                hs.font.italic = True
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(6)
            else:
                hs.font.bold = True
                hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hs.paragraph_format.space_before = Pt(6)
                hs.paragraph_format.space_after = Pt(0)

    return doc


def add_academic_content(doc, markdown_text):
    """Render markdown into an academic-formatted Word doc."""
    font_name = 'Times New Roman'
    font_size = Pt(12)
    table_font_size = Pt(10)

    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    para_lines = []
    in_code_block = False

    def flush_para():
        nonlocal para_lines
        if para_lines:
            text = ' '.join(para_lines)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            apply_inline_formatting(p, text, font_name, font_size)
            para_lines = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        header = [c.strip() for c in table_rows[0].strip('|').split('|')]
        data = []
        for row in table_rows[2:]:
            cells = [c.strip() for c in row.strip('|').split('|')]
            data.append(cells)

        if header and data:
            ncols = len(header)
            tbl = doc.add_table(rows=1 + len(data), cols=ncols)
            tbl.style = 'Table Grid'
            for j, h in enumerate(header):
                if j < ncols:
                    cell = tbl.rows[0].cells[j]
                    cell.text = strip_markdown(h)
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.bold = True
                            r.font.name = font_name
                            r.font.size = table_font_size
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            for ri, rd in enumerate(data):
                for j, ct in enumerate(rd):
                    if j < ncols:
                        cell = tbl.rows[ri + 1].cells[j]
                        cell.text = strip_markdown(ct)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.name = font_name
                                r.font.size = table_font_size
                            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        ptype, level, text = parse_line(line)

        # Handle code blocks
        if ptype == 'code_fence':
            flush_para()
            if not in_code_block:
                in_code_block = True
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        if ptype == 'table_row':
            flush_para()
            if not in_table:
                in_table = True
            table_rows.append(line.strip())
            i += 1
            continue
        elif in_table:
            flush_table()

        if ptype == 'hr':
            flush_para()
        elif ptype == 'heading':
            flush_para()
            if level <= 5:
                h = doc.add_heading(clean_text(text), level=level)
                for r in h.runs:
                    r.font.name = font_name
                    r.font.color.rgb = RGBColor(0, 0, 0)
        elif ptype == 'empty':
            flush_para()
        elif ptype == 'bullet':
            flush_para()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'checklist':
            flush_para()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'numbered':
            flush_para()
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'paragraph':
            para_lines.append(text)
        i += 1

    flush_para()
    if in_table:
        flush_table()


# ============================================================
# Report document (Calibri, colored headers, styled tables)
# ============================================================

def create_report_doc():
    """Create a Word document with report formatting."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    pf = style.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(4)

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    return doc


def add_report_heading(doc, text, level=1):
    """Add a colored heading for report docs."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.name = 'Calibri'
    return h


def add_report_table(doc, headers, data_rows):
    """Add a formatted table with colored header row and alternating stripes."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = strip_markdown(h)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.name = 'Calibri'
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, DARK_BLUE_HEX)

    for ri, rd in enumerate(data_rows):
        for j, ct in enumerate(rd):
            if j < ncols:
                cell = tbl.rows[ri + 1].cells[j]
                cell.text = strip_markdown(str(ct))
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = 'Calibri'
                if ri % 2 == 1:
                    set_cell_shading(cell, ALT_ROW_HEX)

    return tbl


def add_report_content(doc, markdown_text):
    """Render markdown into a report-formatted Word doc."""
    font_name = 'Calibri'
    font_size = Pt(10)

    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    para_lines = []
    in_code_block = False

    def flush_para():
        nonlocal para_lines
        if para_lines:
            text = ' '.join(para_lines)
            p = doc.add_paragraph()
            apply_inline_formatting(p, text, font_name, font_size)
            para_lines = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        header = [c.strip() for c in table_rows[0].strip('|').split('|')]
        data = []
        for row in table_rows[2:]:
            cells = [c.strip() for c in row.strip('|').split('|')]
            data.append(cells)
        if header and data:
            add_report_table(doc, header, data)
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]
        ptype, level, text = parse_line(line)

        # Handle code blocks
        if ptype == 'code_fence':
            flush_para()
            if not in_code_block:
                in_code_block = True
                i += 1
                continue
            else:
                in_code_block = False
                i += 1
                continue
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(line.rstrip())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            i += 1
            continue

        if ptype == 'table_row':
            flush_para()
            if not in_table:
                in_table = True
            table_rows.append(line.strip())
            i += 1
            continue
        elif in_table:
            flush_table()

        if ptype == 'hr':
            flush_para()
        elif ptype == 'heading':
            flush_para()
            add_report_heading(doc, clean_text(text), level=min(level, 4))
        elif ptype == 'empty':
            flush_para()
        elif ptype == 'bullet':
            flush_para()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'checklist':
            flush_para()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'numbered':
            flush_para()
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text, font_name, font_size)
        elif ptype == 'paragraph':
            para_lines.append(text)
        i += 1

    flush_para()
    if in_table:
        flush_table()


# ============================================================
# Document builders
# ============================================================

def build_memo():
    """Build Deliverable 1: Qualitative Pivot Justification Memo."""
    print("Building: Qualitative Pivot Justification Memo...")
    doc = create_academic_doc()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Memorandum")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Read markdown and add content
    md_path = DELIVERABLES_DIR / '01_qualitative_pivot_memo.md'
    md_text = md_path.read_text()
    add_academic_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_Qualitative_Pivot_Memo.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def build_journals():
    """Build Deliverable 2: Target Journal Strategy."""
    print("Building: Target Journal Strategy...")
    doc = create_report_doc()

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Target Journal Strategy")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quantitative Source Selection Paper")
    run.font.size = Pt(16)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for Richard Davidson\nMarch 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_page_break()

    md_path = DELIVERABLES_DIR / '02_target_journals.md'
    md_text = md_path.read_text()
    add_report_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_Target_Journals.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def build_case_study():
    """Build Deliverable 3: Case Study Research Design."""
    print("Building: Case Study Research Design...")
    doc = create_academic_doc()

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    run = p.add_run(
        "Qualitative Case Study Research Design:\n"
        "Understanding Source Selection Decision-Making\n"
        "in Federal Procurement"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Richard Davidson\nDBA Candidate, Kelley School of Business\nIndiana University")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("March 2026")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    md_path = DELIVERABLES_DIR / '03_case_study_design.md'
    md_text = md_path.read_text()
    add_academic_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_Case_Study_Design.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def build_software_guide():
    """Build Deliverable 4: Qualitative Software Guide."""
    print("Building: Qualitative Software Guide...")
    doc = create_report_doc()

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Qualitative Analysis\nSoftware Guide")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pre-Program Learning Plan for the Kelley DBA")
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared for Richard Davidson\nMarch 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_page_break()

    md_path = DELIVERABLES_DIR / '04_qualitative_software_guide.md'
    md_text = md_path.read_text()
    add_report_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_Qualitative_Software_Guide.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def build_revised_proposal():
    """Build Deliverable 5: Revised Research Proposal V3."""
    print("Building: Revised Research Proposal V3...")
    doc = create_academic_doc()

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "When Does Method Matter?\n"
        "Source Selection Decision-Making and\n"
        "Contract Outcomes in Federal Procurement"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A DBA Research Proposal (Revised)")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run.italic = True

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Richard Davidson\n"
        "Indiana University\n"
        "Kelley School of Business, DBA Program\n"
        "March 2026"
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    doc.add_page_break()

    md_path = DELIVERABLES_DIR / '05_revised_research_proposal.md'
    md_text = md_path.read_text()
    add_academic_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_Research_Proposal_V3.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


def build_arp_plan():
    """Build Deliverable 6: ARP I-IV Plan."""
    print("Building: ARP I-IV Plan...")
    doc = create_report_doc()

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Applied Research Practicum\n(ARP) Plan")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Source Selection Research\nARP I-IV Pipeline to Dissertation")
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM_BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Richard Davidson\nKelley School of Business, DBA Program\nMarch 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_page_break()

    md_path = DELIVERABLES_DIR / '06_arp_plan.md'
    md_text = md_path.read_text()
    add_report_content(doc, md_text)

    out_path = OUTPUT_DIR / 'Davidson_ARP_Plan.docx'
    doc.save(str(out_path))
    print(f"  Saved: {out_path}")
    return out_path


# ============================================================
# Main
# ============================================================

def main():
    print("=" * 60)
    print("BUILDING PhD PIVOT DELIVERABLES")
    print("=" * 60)
    print()

    paths = []
    paths.append(build_memo())
    paths.append(build_journals())
    paths.append(build_case_study())
    paths.append(build_software_guide())
    paths.append(build_revised_proposal())
    paths.append(build_arp_plan())

    print()
    print("=" * 60)
    print("ALL DELIVERABLES BUILT SUCCESSFULLY")
    print("=" * 60)
    print()
    for p in paths:
        print(f"  {p.name}")
    print()
    print(f"Output directory: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
