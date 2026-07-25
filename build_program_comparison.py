#!/usr/bin/env python3
"""Build Executive PhD/DBA Programs Comparison Word Document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2E4057')

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)
    return h

def add_program(doc, name, details):
    """Add a program entry with name and detail bullets."""
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x8A)

    for key, val in details:
        p = doc.add_paragraph(style='List Bullet')
        run_key = p.add_run(f"{key}: ")
        run_key.bold = True
        run_key.font.size = Pt(9)
        run_val = p.add_run(str(val))
        run_val.font.size = Pt(9)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Executive PhD & DBA Programs\nGlobal Comparison Guide")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Prepared for Richard Davidson\nFebruary 2026")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
criteria = doc.add_paragraph()
criteria.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = criteria.add_run(
    "Criteria: PhD or DBA | No GRE/GMAT preferred | Remote/online preferred\n"
    "AACSB/EQUIS/AMBA accredited | Public procurement/supply chain flagged"
)
run.font.size = Pt(11)
run.font.italic = True

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading(doc, "Table of Contents", level=1)
toc_items = [
    "1. Overall Top 10 Recommendations",
    "2. Programs with Strongest Procurement/Supply Chain Alignment",
    "3. United States Programs",
    "4. Europe Programs",
    "5. Asia-Pacific Programs",
    "6. Middle East Programs",
    "7. US Programs without AACSB (Reference)",
    "8. Programs Requiring GMAT (Reference)",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ============================================================
# 1. OVERALL TOP 10
# ============================================================
add_heading(doc, "1. Overall Top 10 Recommendations", level=1)
p = doc.add_paragraph(
    "Weighing all criteria: no GRE/GMAT, remote delivery, accreditation, cost, and procurement relevance."
)
p.italic = True

add_table(doc,
    ["#", "Program", "Cost (USD)", "Format", "Accreditation", "Key Advantage"],
    [
        ["1", "Texas State (McCoy) DBA", "$36,900", "100% online", "AACSB", "Lowest-cost AACSB, fully async"],
        ["2", "Edinburgh/Heriot-Watt DBA", "~$19-32K", "100% online", "Triple Crown", "Most flexible, most affordable"],
        ["3", "Aston Business School DBA", "~$50K", "100% online", "Triple Crown", "Best all-criteria match globally"],
        ["4", "Indiana Kelley EDBA", "$96-159K", "Near-remote", "AACSB", "Top brand, only 2 residencies"],
        ["5", "Grenoble GEM DBA", "~$20-45K", "Blended", "Triple Crown", "3-year fast track"],
        ["6", "Cranfield DBA", "~$75K", "Blended", "Triple Crown", "Best procurement/SCM faculty"],
        ["7", "Rutgers Asia Pacific DBA", "Contact", "Blended", "AACSB", "World-class SCM program"],
        ["8", "Aston/Kaplan Singapore DBA", "~$45K", "Blended", "Triple Crown", "Great Asia-Pac value"],
        ["9", "USM Malaysia DBA", "~$17K", "Mixed/online", "AACSB", "Ultra-affordable"],
        ["10", "UAE University DBA", "~$41-55K", "Weekend", "AACSB", "Pioneer Gulf DBA"],
    ],
    col_widths=[0.3, 2.0, 1.0, 1.0, 1.1, 2.0]
)

doc.add_paragraph()
add_heading(doc, "Best Value + Fully Online", level=3)
doc.add_paragraph("Texas State, Edinburgh/Heriot-Watt, and Aston are the clear winners for fully online, no GRE, and top accreditation. Range: $19K to $50K total.")

add_heading(doc, "Best Prestige + Mostly Remote", level=3)
doc.add_paragraph("Indiana Kelley EDBA and Grenoble GEM DBA offer top-tier brand recognition with minimal in-person requirements.")

doc.add_page_break()

# ============================================================
# 2. PROCUREMENT/SUPPLY CHAIN ALIGNMENT
# ============================================================
add_heading(doc, "2. Programs with Strongest Procurement/Supply Chain Alignment", level=1)

add_table(doc,
    ["Program", "Region", "Why It's Relevant"],
    [
        ["Cranfield (UK)", "Europe", "Top 12 globally for SCM; dedicated Procurement & Supply Chain MSc; faculty supervisors in procurement"],
        ["SDA Bocconi (Italy)", "Europe", "Executive Programs in Procurement & Supply Management; Italian public procurement research"],
        ["ASU W.P. Carey (US)", "US", "Only AACSB exec DBA with formal Supply Chain Management specialization (GMAT required)"],
        ["Rutgers Asia Pacific (Singapore)", "Asia", "Globally renowned SCM program; specialization possible in DBA"],
        ["Victoria Univ. Wellington (NZ)", "Asia-Pac", "School of Business AND Government; explicit government procurement research"],
        ["PolyU Hong Kong", "Asia", "One of Asia's strongest Logistics/SCM departments"],
        ["Paris-Dauphine (France)", "Europe", "Strong public economics, public policy; procurement/public value research"],
        ["Maastricht MSM (Netherlands)", "Europe", "Development economics, governance, public sector management focus"],
        ["Liberty University (US)", "US", "Only fully online DBA with Supply Chain/Logistics track (ACBSP, not AACSB)"],
        ["Qatar University", "Middle East", "Qatar's massive procurement reform agenda (GMAT required)"],
    ],
    col_widths=[2.0, 0.8, 4.5]
)

doc.add_page_break()

# ============================================================
# 3. UNITED STATES
# ============================================================
add_heading(doc, "3. United States Programs", level=1)

# --- Tier 1 ---
add_heading(doc, "Tier 1: Best Fit (AACSB + No GRE/GMAT + Online/Hybrid)", level=2)

add_program(doc, "1. Texas State University (McCoy College) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "100% fully online, asynchronous. No residency required."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "$36,900 total (flat-rate, 5 payments/year)"),
    ("Duration", "3 years (cohort-based)"),
    ("Notes", "Launched Fall 2025. Lowest-cost AACSB program found."),
    ("Rating", "EXCELLENT"),
])

add_program(doc, "2. Marshall University (Brad D. Smith Schools) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. 5 on-campus weekend residencies/year (Fri-Sat) Years 1-2, otherwise online."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$850/credit hour (~$51K estimated total)"),
    ("Duration", "3 years"),
    ("Notes", "WRDS/Compustat/CRSP database access included. Member of Executive DBA Council."),
    ("Rating", "STRONG"),
])

add_program(doc, "3. Drexel University (LeBow College) -- Executive DBA", [
    ("Accreditation", "AACSB (STEM-designated)"),
    ("Format", "Blended. Two face-to-face residencies per 10-week quarter (Thu-Sat, Philadelphia)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$128,000 total"),
    ("Duration", "2.5 years (accelerated)"),
    ("Notes", "R1 research institution. Supply chain faculty present."),
    ("Rating", "STRONG (frequent residencies and high cost)"),
])

add_program(doc, "4. University of South Florida (Muma College) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Online (alternate Saturdays, synchronous) OR in-person (Tampa)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "$90,000 total ($30K/year)"),
    ("Duration", "3 years"),
    ("Notes", "Requires 12 years professional + 5 at senior level."),
    ("Rating", "GOOD"),
])

add_program(doc, "5. Indiana University (Kelley School) -- Executive DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Predominantly online. Two required in-person residencies total."),
    ("GRE/GMAT", "OPTIONAL"),
    ("Tuition", "$96,600-$158,700 (depends on prior MBA credits)"),
    ("Duration", "~3 years"),
    ("Notes", "Top-ranked online business programs. Strong brand recognition."),
    ("Rating", "GOOD"),
])

add_program(doc, "6. University of Wisconsin-Whitewater -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Primarily online; one mandatory on-campus Saturday in May."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$96,000 total (same rate all students)"),
    ("Duration", "3 years"),
    ("Rating", "GOOD"),
])

add_program(doc, "7. University of Michigan-Flint -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "100% online, part-time"),
    ("GRE/GMAT", "Not mentioned (verify directly)"),
    ("Tuition", "~$80,000 total (tuition locked for cohort)"),
    ("Duration", "3 years"),
    ("Rating", "GOOD"),
])

add_program(doc, "8. Florida Atlantic University -- Executive PhD", [
    ("Accreditation", "AACSB"),
    ("Format", "On-campus Saturdays (Boca Raton, FL)"),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$81,000 total"),
    ("Duration", "3-4 years"),
    ("Notes", "One of the only AACSB schools offering PhD (not DBA) in executive format."),
    ("Rating", "GOOD (primarily in-person)"),
])

add_program(doc, "9. University of Missouri-St. Louis -- DBA", [
    ("Accreditation", "AACSB (dual-accredited business + accounting)"),
    ("Format", "Hybrid. Two weekend visits to campus per semester."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "$105,840 (auto $10K scholarship = ~$95,840 net)"),
    ("Duration", "3 years"),
    ("Rating", "GOOD"),
])

# --- Tier 2 ---
add_heading(doc, "Tier 2: Solid Programs with Some Caveats", level=2)

add_program(doc, "10. Temple University (Fox School) -- Executive DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. Six on-site weekend residencies/year (Fri-Sun) in Philadelphia, Years 1-2."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$118,800 total (includes residency costs, books, lodging)"),
    ("Duration", "3 years"),
])

add_program(doc, "11. Kennesaw State University (Coles College) -- PhD", [
    ("Accreditation", "AACSB (STEM-designated)"),
    ("Format", "Monthly Fri-Sun residencies on campus (Kennesaw, GA) Years 1-2."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$34,000+ (in-state rates)"),
    ("Duration", "3+ years"),
    ("Notes", "Rare PhD (not DBA) in executive format with no GRE. Requires 10+ years experience."),
])

add_program(doc, "12. Oklahoma State University (Spears School) -- Executive PhD & DBA", [
    ("Accreditation", "AACSB (dual-accredited)"),
    ("Format", "Hybrid. Monthly 3-day residencies in Stillwater, OK."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "PhD: $139,500 | DBA: $129,000"),
    ("Duration", "PhD: 3 years | DBA: 3.5 years"),
])

add_program(doc, "13. Case Western Reserve (Weatherhead) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. Three on-campus + three virtual residencies per semester (Cleveland)."),
    ("GRE/GMAT", "WAIVABLE (5+ yrs experience, 3.0+ GPA, or CPA/CFA)"),
    ("Tuition", "~$54,400/year (~$163K total)"),
    ("Duration", "3 years"),
])

add_program(doc, "14. Penn State University (Smeal College) -- Executive DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. 4-5 days in-person at start of each semester, then alternating weekend remote."),
    ("GRE/GMAT", "Likely NOT REQUIRED (verify directly)"),
    ("Tuition", "$138,000 total (includes residency costs)"),
    ("Duration", "~3 years"),
])

add_program(doc, "15. Georgia State University (Robinson College) -- Executive DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "In-person 3-day residencies monthly (Atlanta)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "$115,313-$129,497 (all-inclusive)"),
    ("Duration", "3 years"),
])

add_program(doc, "16. Creighton University (Heider College) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. 7 in-person residencies/year Years 1-2 (all-inclusive incl. hotel/meals)."),
    ("GRE/GMAT", "OPTIONAL"),
    ("Tuition", "All-inclusive; contact program"),
    ("Duration", "3 years"),
])

add_program(doc, "17. Prairie View A&M University -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Hybrid. Monthly two-day weekend residencies (Texas Triangle)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$8K/yr in-state; ~$17K/yr out-of-state"),
    ("Duration", "3 years (60 credits)"),
])

add_program(doc, "18. University of Dallas (Gupta College) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "Blended. One weekend/month on campus (Irving, TX)."),
    ("GRE/GMAT", "Likely NOT REQUIRED (verify)"),
    ("Tuition", "Not published"),
    ("Duration", "~3 years"),
])

add_program(doc, "19. UT Dallas (Naveen Jindal School) -- DBA", [
    ("Accreditation", "AACSB"),
    ("Format", "In-person, three Saturdays/month (Richardson, TX)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "$120,000 total"),
    ("Duration", "3 years"),
])

add_heading(doc, "Notable: GRE Required but Topically Relevant", level=3)
add_program(doc, "ASU W.P. Carey -- DBA in Supply Chain Management", [
    ("Accreditation", "AACSB"),
    ("GRE/GMAT", "REQUIRED (no waivers)"),
    ("Notes", "The ONLY AACSB executive DBA with a formal Supply Chain Management specialization. Most directly aligned with your dissertation topic."),
])

doc.add_page_break()

# ============================================================
# 4. EUROPE
# ============================================================
add_heading(doc, "4. Europe Programs", level=1)

add_heading(doc, "Tier 1: Best Fit (Triple Crown + No GMAT + Online/Hybrid)", level=2)

add_program(doc, "1. Aston Business School -- Executive DBA (Online)", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Fully online, no compulsory live lectures, no fixed schedule."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "GBP 39,000 total (~$50K); monthly payment options"),
    ("Duration", "4-6 years"),
    ("Notes", "Perhaps the strongest match on ALL criteria globally."),
    ("Rating", "EXCELLENT"),
])

add_program(doc, "2. Edinburgh Business School / Heriot-Watt University -- DBA", [
    ("Country", "UK (Scotland) -- globally accessible online"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Fully online, self-study distance learning. No residencies. No fixed start date."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Per-module pricing; estimated GBP 15,000-25,000 total (~$19K-32K)"),
    ("Duration", "~5 years average"),
    ("Notes", "25+ years of online delivery. Most flexible option in Europe. Also has Dubai campus."),
    ("Rating", "EXCELLENT"),
])

add_program(doc, "3. University of Liverpool Management School -- Executive DBA", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Hybrid -- primarily online with 4 weekend on-campus residencies during first 2 years."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "See course page; 15% discount available"),
    ("Duration", "4.5 years standard; up to 7.5 years"),
    ("Notes", "Action-research model fits applied procurement dissertation work well."),
    ("Rating", "STRONG"),
])

add_program(doc, "4. Grenoble Ecole de Management (GEM) -- Global DBA", [
    ("Country", "France"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Blended -- online workshops + ~3 in-person residential seminars/year."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Merit scholarships up to EUR 6,000; contact program"),
    ("Duration", "3 years"),
    ("Notes", "First DBA program in Europe. Supply chain and sustainability are active research areas."),
    ("Rating", "STRONG"),
])

add_program(doc, "5. Nottingham Business School (NTU) -- DBA", [
    ("Country", "UK"),
    ("Accreditation", "AACSB + EQUIS (double accreditation)"),
    ("Format", "Blended -- distance/online with some in-person components."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Contact school; generally more affordable than Russell Group"),
    ("Duration", "4-6 years"),
    ("Notes", "One of the oldest DBA programs in the UK (20+ years)."),
])

add_heading(doc, "Tier 2: Strong Programs", level=2)

add_program(doc, "6. Cranfield School of Management -- Executive DBA  [PROCUREMENT/SCM FLAGGED]", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Blended -- 8 residential weeks over 4 years (four 5-day modules Yrs 1-2, then supervision)."),
    ("GRE/GMAT", "Not explicitly required (verify)"),
    ("Tuition", "GBP 18,000/yr (Yrs 1-2), GBP 11,500/yr (Yrs 3-4) = ~GBP 59,000 total (~$75K)"),
    ("Duration", "4 years"),
    ("Procurement/SC", "TOP 12 GLOBALLY for supply chain management. Dedicated MSc in Procurement & Supply Chain Management. Faculty supervisors with procurement expertise."),
    ("Rating", "STRONG -- Best European supply chain/procurement research environment"),
])

add_program(doc, "7. Henley Business School (University of Reading) -- DBA", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown; DBA itself holds AMBA accreditation (1 of only 11 DBAs globally)"),
    ("Format", "Part-time blended with residential + distance components."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "GBP 13,900/year = ~GBP 55,600 total (~$71K)"),
    ("Duration", "4 years"),
])

add_program(doc, "8. Durham University Business School -- DBA", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Part-time blended; residentials in Durham; distance supervision."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "GBP 50,000 total (~$64K)"),
    ("Duration", "4-6 years"),
])

add_program(doc, "9. Alliance Manchester Business School -- DBA", [
    ("Country", "UK"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Blended -- two 4-5 day residentials/year + online coursework."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Contact school (~GBP 15-20K/year estimated)"),
    ("Duration", "4 years"),
])

add_program(doc, "10. SDA Bocconi School of Management -- DBA  [PROCUREMENT FLAGGED]", [
    ("Country", "Italy"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Hybrid -- 4 week-long residential modules in Milan + 4 virtual modules."),
    ("GRE/GMAT", "NOT REQUIRED (requires 15+ years experience)"),
    ("Tuition", "~EUR 90,000 total; merit scholarships up to EUR 32,000"),
    ("Duration", "3 years"),
    ("Procurement/SC", "Dedicated Executive Programs in Procurement & Supply Management. Italian public procurement is a recognized research focus."),
])

add_program(doc, "11. Paris-Dauphine University (PSL) -- Executive PhD  [PUBLIC ADMIN FLAGGED]", [
    ("Country", "France"),
    ("Accreditation", "EQUIS"),
    ("Format", "Hybrid -- 1 week/term at Paris campus (3x/year) + online 2 days/month."),
    ("GRE/GMAT", "Not mentioned"),
    ("Tuition", "Contact school"),
    ("Duration", "3-5 years"),
    ("Procurement/SC", "Strong public economics, public policy research. Procurement and public value fits well."),
])

add_program(doc, "12. Maastricht School of Management -- Executive PhD  [PUBLIC SECTOR FLAGGED]", [
    ("Country", "Netherlands"),
    ("Accreditation", "Triple Crown via Maastricht University"),
    ("Format", "Blended -- on-campus workshops + online learning."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Contact school"),
    ("Duration", "4-7 years"),
    ("Procurement/SC", "Focus on development economics, governance, and public sector management."),
])

add_program(doc, "13. RSM Erasmus University -- Part-Time PhD in Management", [
    ("Country", "Netherlands"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Part-time -- eight modules of 3-4 consecutive days in Rotterdam."),
    ("GRE/GMAT", "OPTIONAL"),
    ("Tuition", "~EUR 46,000 max total (~$50K)"),
    ("Duration", "5-6 years"),
    ("Notes", "True academic PhD (not DBA). Strong supply chain research faculty."),
])

add_program(doc, "14. Vlerick Business School -- Executive PhD", [
    ("Country", "Belgium"),
    ("Accreditation", "EQUIS + AMBA (joint degree from KU Leuven and Ghent University)"),
    ("Format", "Blended -- online platform + on-campus acceleration weeks."),
    ("GRE/GMAT", "Waivable (Vlerick VBAT assessment alternative)"),
    ("Tuition", "EUR 66,000 total (~$72K)"),
    ("Duration", "5 years"),
])

doc.add_page_break()

# ============================================================
# 5. ASIA-PACIFIC
# ============================================================
add_heading(doc, "5. Asia-Pacific Programs", level=1)

add_heading(doc, "Tier 1: Best Fit", level=2)

add_program(doc, "1. Aston University / Kaplan Singapore -- Executive DBA", [
    ("Country", "Singapore (UK-awarding institution)"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Blended -- face-to-face taught modules at Kaplan Singapore; online supervision for research."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "SGD 59,950 after grant (~$45,000 USD)"),
    ("Duration", "4 years (max 6)"),
    ("Rating", "EXCELLENT"),
])

add_program(doc, "2. Rutgers Business School Asia Pacific -- DBA  [SUPPLY CHAIN FLAGGED]", [
    ("Country", "Singapore"),
    ("Accreditation", "AACSB"),
    ("Format", "Blended/hybrid -- core modules in-person weekends/evenings; online research supervision."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Contact directly"),
    ("Duration", "2+ years"),
    ("Procurement/SC", "Rutgers is globally renowned for supply chain management. SCM specialization possible."),
    ("Rating", "STRONG"),
])

add_program(doc, "3. City University of Hong Kong -- DBA", [
    ("Country", "Hong Kong"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA) -- ONLY triple-crown in Greater China"),
    ("Format", "Hybrid weekend mode (Sat afternoon + Sun) once/month + 2 residential workshops."),
    ("GRE/GMAT", "NOT REQUIRED (waived with 2+ years experience)"),
    ("Tuition", "~$121,500 USD"),
    ("Duration", "3-6 years"),
])

add_program(doc, "4. Sasin / Chulalongkorn University -- DBA", [
    ("Country", "Thailand"),
    ("Accreditation", "AACSB + EQUIS (co-founded with Kellogg and Wharton)"),
    ("Format", "Hybrid -- ~14 in-person intensive days/year at Bangkok + online."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$112,000 USD (includes 2 international conferences)"),
    ("Duration", "3 years"),
    ("Notes", "Requires 10+ years management experience."),
])

add_heading(doc, "Tier 2: Strong Options", level=2)

add_program(doc, "5. Hong Kong Polytechnic University -- DBA  [SUPPLY CHAIN FLAGGED]", [
    ("Country", "Hong Kong"),
    ("Accreditation", "AACSB + EQUIS"),
    ("Format", "In-person, part-time weekend/evening. Collaboration with IMD + LSE."),
    ("GRE/GMAT", "Not explicitly required"),
    ("Tuition", "~$202,000 USD"),
    ("Duration", "2-6 years"),
    ("Procurement/SC", "One of Asia's strongest Logistics/SCM departments."),
])

add_program(doc, "6. HKUST -- DBA", [
    ("Country", "Hong Kong"),
    ("Accreditation", "AACSB + EQUIS"),
    ("Format", "In-person modular -- 4 consecutive days per module."),
    ("GRE/GMAT", "Not explicitly required for DBA"),
    ("Tuition", "~$230,000 USD (includes accommodations)"),
    ("Duration", "4 years"),
])

add_program(doc, "7. Hong Kong Baptist University -- DBA", [
    ("Country", "Hong Kong"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Part-time, in-person."),
    ("GRE/GMAT", "WAIVED with 8+ years experience"),
    ("Tuition", "~$101,000 USD"),
    ("Duration", "3-4 years"),
])

add_program(doc, "8. Singapore Management University -- DBA", [
    ("Country", "Singapore"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "In-person modular 3-day residentials (Fri-Sun)."),
    ("GRE/GMAT", "Not explicitly stated"),
    ("Tuition", "Contact directly"),
    ("Duration", "3 years minimum"),
])

add_program(doc, "9. Universiti Sains Malaysia (USM) -- DBA", [
    ("Country", "Malaysia"),
    ("Accreditation", "AACSB"),
    ("Format", "Mixed mode delivery / online mode available."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$17,475 USD total (one of the most affordable globally)"),
    ("Duration", "TBD"),
])

add_program(doc, "10. Universiti Malaya -- DBA", [
    ("Country", "Malaysia"),
    ("Accreditation", "AACSB + AMBA"),
    ("Format", "In-person weekend classes (Sat/Sun)."),
    ("GRE/GMAT", "Not specified"),
    ("Tuition", "~$17,500 USD (international)"),
    ("Duration", "3-6 years"),
])

add_heading(doc, "Australia & New Zealand", level=2)

add_program(doc, "11. Deakin University -- DBA", [
    ("Country", "Australia"),
    ("Accreditation", "AACSB + EQUIS"),
    ("Format", "Research-based; 50+ years of distance/online expertise -- likely distance-compatible."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Verify; domestic may be RTP-funded"),
    ("Duration", "4-6 years part-time"),
])

add_program(doc, "12. Curtin University -- DBA", [
    ("Country", "Australia (Perth)"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "In-person Perth; research-based with supervisor."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "FREE for Australian citizens/PR/NZ (RTP-funded). International: contact."),
    ("Duration", "2-6 years"),
])

add_program(doc, "13. Victoria University of Wellington -- PhD  [GOVERNMENT/PROCUREMENT FLAGGED]", [
    ("Country", "New Zealand"),
    ("Accreditation", "Triple Crown (AACSB + EQUIS + AMBA)"),
    ("Format", "Research doctorate; primarily campus-based."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$6,000/year"),
    ("Procurement/SC", "School of Business AND Government. Explicit government procurement research."),
])

add_program(doc, "14. Massey University -- DBA", [
    ("Country", "New Zealand"),
    ("Accreditation", "AACSB"),
    ("Format", "Mixed/distance (60 years of distance education)."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "Contact school"),
    ("Duration", "3-6 years"),
])

add_heading(doc, "Notable: GMAT Required but Exceptional Value", level=3)
add_program(doc, "Hitotsubashi ICS (Japan) -- DBA", [
    ("Accreditation", "AACSB"),
    ("GRE/GMAT", "REQUIRED (GMAT 600+ or GRE 160Q/150V)"),
    ("Tuition", "~$25,000 USD total (exceptionally affordable)"),
    ("Duration", "3 years"),
])

doc.add_page_break()

# ============================================================
# 6. MIDDLE EAST
# ============================================================
add_heading(doc, "6. Middle East Programs", level=1)

add_heading(doc, "Tier 1: Best Fit", level=2)

add_program(doc, "1. Hult International Business School -- DBA", [
    ("Country", "Dubai + rotating (Boston, London, San Francisco, Dubai)"),
    ("Accreditation", "Triple Crown (AACSB + AMBA + EQUIS)"),
    ("Format", "Hybrid -- weekly live Zoom (1-2x/week) + three 5-day residencies/year rotating globally."),
    ("GRE/GMAT", 'NOT REQUIRED ("no specific GRE, GMAT or GPA requirements")'),
    ("Tuition", "$119,000 USD total (travel to residencies additional)"),
    ("Duration", "3 years"),
    ("Notes", "Requires 8-10 years total experience, 5 at managerial/executive level."),
    ("Rating", "STRONG"),
])

add_program(doc, "2. UAE University (UAEU) -- DBA", [
    ("Country", "UAE (Al Ain)"),
    ("Accreditation", "AACSB"),
    ("Format", "Part-time, weekend in-person."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$41K-55K USD estimated"),
    ("Duration", "3-4 years"),
    ("Notes", "Pioneer AACSB-accredited DBA in the UAE."),
    ("Rating", "STRONG"),
])

add_program(doc, "3. Abu Dhabi University -- DBA", [
    ("Country", "UAE (Abu Dhabi / Dubai campuses)"),
    ("Accreditation", "AACSB + EQUIS (first dual-accredited in UAE)"),
    ("Format", "Part-time; confirm hybrid options."),
    ("GRE/GMAT", "Not listed in requirements"),
    ("Tuition", "~$52K-65K USD estimated"),
    ("Duration", "4-7 years (66 credit hours)"),
    ("Rating", "STRONG (dual accreditation exceptional for region)"),
])

add_program(doc, "4. Edinburgh Business School / Heriot-Watt (Dubai campus) -- DBA", [
    ("Country", "UK-based, Dubai campus available"),
    ("Accreditation", "AACSB"),
    ("Format", "100% online globally; optional study at Dubai campus."),
    ("GRE/GMAT", "NOT REQUIRED"),
    ("Tuition", "~$19K-32K USD total (most affordable accredited option)"),
    ("Duration", "~5 years"),
    ("Notes", "Also listed under Europe."),
])

add_heading(doc, "Tier 2: Viable Options", level=2)

add_program(doc, "5. Ajman University -- DBA", [
    ("Country", "UAE (Ajman)"),
    ("Accreditation", "AACSB"),
    ("Format", "Weekend in-person (Sat/Sun)."),
    ("GRE/GMAT", "Not listed in requirements"),
    ("Tuition", "~$73,500 USD; scholarships available (alumni: 20%)"),
    ("Duration", "3-4 years"),
])

add_program(doc, "6. SP Jain School of Global Management (Dubai) -- DBA", [
    ("Country", "Dubai"),
    ("Accreditation", "AACSB + AMBA (verify scope)"),
    ("Format", "100% online option available; also on-campus Dubai/Sydney."),
    ("GRE/GMAT", "WAIVABLE for strong candidates"),
    ("Tuition", "~$28K-30K USD (online version)"),
    ("Duration", "3-5 years"),
])

add_program(doc, "7. University of Wollongong Dubai -- DBA", [
    ("Country", "UAE (Dubai)"),
    ("Accreditation", "AACSB (via UOW Australia); first CAA-accredited DBA in UAE"),
    ("Format", "In-person Dubai campus; very small cohort (10 students/intake)."),
    ("GRE/GMAT", "Unconfirmed -- inquire directly"),
    ("Tuition", "~$45K-52K USD total"),
    ("Duration", "3-4 years"),
    ("Notes", "Has MSc in Logistics and Supply Chain Management."),
])

add_program(doc, "8. University of Sharjah -- DBA", [
    ("Country", "UAE (Sharjah)"),
    ("Accreditation", "AACSB"),
    ("Format", "Contact directly"),
    ("GRE/GMAT", "Contact directly"),
])

doc.add_page_break()

# ============================================================
# 7. US PROGRAMS WITHOUT AACSB
# ============================================================
add_heading(doc, "7. US Programs without AACSB (Reference)", level=1)
doc.add_paragraph("These programs do not carry AACSB, EQUIS, or AMBA accreditation but may be of interest for cost or format reasons.")

add_table(doc,
    ["Program", "Accreditation", "GRE/GMAT", "Format", "Est. Cost", "Notes"],
    [
        ["Wilmington University", "IACBE", "None", "100% online", "~$25,776", "Lowest cost found"],
        ["UMGC", "IACBE", "None", "Hybrid (3 MD res/yr)", "~$52,176", ""],
        ["Liberty University", "ACBSP", "None", "100% online", "~$22,500", "Supply Chain/Logistics DBA track"],
        ["Capella University", "ACBSP", "None", "100% online", "Capped $45K", ""],
        ["Walden University", "ACBSP", "None", "Online + 2 residencies", "~$79,477", ""],
        ["National University", "ACBSP", "None", "100% online", "Per credit", ""],
    ],
    col_widths=[1.5, 0.8, 0.6, 1.3, 0.9, 2.0]
)

doc.add_paragraph()

# ============================================================
# 8. PROGRAMS REQUIRING GMAT
# ============================================================
add_heading(doc, "8. Programs Requiring GMAT/GRE (Reference)", level=1)
doc.add_paragraph("These programs require standardized testing but may be worth considering for their topical relevance or exceptional value.")

add_table(doc,
    ["Program", "Country", "Accreditation", "Notes"],
    [
        ["ASU W.P. Carey DBA (Supply Chain)", "US", "AACSB", "Only formal SCM DBA; directly aligned with dissertation"],
        ["Hitotsubashi ICS DBA", "Japan", "AACSB", "~$25K total; exceptional value"],
        ["American University of Sharjah PhD", "UAE", "AACSB", "Academic PhD; GRE/GMAT required"],
        ["Qatar University PhD", "Qatar", "AACSB", "GMAT 400+; great procurement policy context"],
        ["Koc University PhD", "Turkey", "Triple Crown", "Only triple-crown in Turkey; full-time"],
        ["Sabanci University PhD", "Turkey", "AACSB", "Full-time in-person"],
        ["Bilkent University PhD", "Turkey", "AACSB", "First AACSB-accredited in Turkey"],
        ["Pace University DPS", "US", "AACSB", "Oldest US exec doctoral program (1972)"],
    ],
    col_widths=[2.5, 0.9, 1.1, 3.0]
)

doc.add_paragraph()
doc.add_paragraph()

# --- Disclaimer ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Note: All information researched February 2026. Tuition, formats, and admission requirements "
    "change regularly. Verify all details directly with programs before applying."
)
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.expanduser("~/Claude/phd-research/Executive_PhD_DBA_Programs_Global_Comparison.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
