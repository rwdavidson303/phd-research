#!/usr/bin/env python3
"""
Dissertation Document Builder
Assembles all Markdown chapters into a formatted Word document
conforming to Indiana University formatting requirements.

IU Requirements:
- 12pt Times New Roman throughout
- 1.5" left margin, 1" top/right/bottom (2" top on chapter openers)
- Double-spaced text; single-spaced block quotes
- Front matter: lowercase Roman numerals (ii, iii...)
- Body: Arabic numerals (1, 2, 3...)
- APA-style headings (Level 1-5)
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT = Path(__file__).parent.parent
CHAPTERS_DIR = Path(__file__).parent / 'chapters'
OUTPUT_DIR = Path(__file__).parent / 'output'
FIGURES_DIR = PROJECT / 'analysis' / 'output' / 'figures'
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def create_document():
    """Create a new Word document with IU formatting."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set paragraph formatting
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Set margins for first section
    section = doc.sections[0]
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Configure heading styles
    for level in range(1, 6):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = 'Times New Roman'
            h_font.color.rgb = RGBColor(0, 0, 0)

            if level == 1:
                h_font.size = Pt(12)
                h_font.bold = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                h_style.paragraph_format.space_before = Pt(24)
                h_style.paragraph_format.space_after = Pt(12)
            elif level == 2:
                h_font.size = Pt(12)
                h_font.bold = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h_style.paragraph_format.space_before = Pt(12)
                h_style.paragraph_format.space_after = Pt(6)
            elif level == 3:
                h_font.size = Pt(12)
                h_font.bold = True
                h_font.italic = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h_style.paragraph_format.space_before = Pt(12)
                h_style.paragraph_format.space_after = Pt(6)
            elif level == 4:
                h_font.size = Pt(12)
                h_font.bold = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h_style.paragraph_format.first_line_indent = Inches(0.5)
                h_style.paragraph_format.space_before = Pt(6)
                h_style.paragraph_format.space_after = Pt(0)
            elif level == 5:
                h_font.size = Pt(12)
                h_font.bold = True
                h_font.italic = True
                h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                h_style.paragraph_format.first_line_indent = Inches(0.5)
                h_style.paragraph_format.space_before = Pt(6)
                h_style.paragraph_format.space_after = Pt(0)

            h_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    return doc


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def clean_html_entities(text):
    """Replace HTML entities with their plain text equivalents."""
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&mdash;', '\u2014')
    text = text.replace('&ndash;', '\u2013')
    text = text.replace('---', '\u2014')
    text = text.replace('--', '\u2013')
    return text


def strip_markdown_formatting(text):
    """Remove markdown bold/italic markers and return plain text."""
    text = clean_html_entities(text)
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


def parse_markdown_line(line):
    """Determine the type and content of a markdown line."""
    stripped = line.strip()

    # Horizontal rule - skip
    if stripped in ('---', '***', '___'):
        return ('hr', 0, '')

    # Heading levels
    if stripped.startswith('#'):
        level = 0
        for ch in stripped:
            if ch == '#':
                level += 1
            else:
                break
        text = stripped[level:].strip()
        return ('heading', level, text)

    # Empty line
    if not stripped:
        return ('empty', 0, '')

    # Table row
    if stripped.startswith('|') and stripped.endswith('|'):
        return ('table_row', 0, stripped)

    # List item
    if stripped.startswith('- ') or stripped.startswith('* '):
        return ('list_item', 0, stripped[2:])

    # Numbered list
    match = re.match(r'^(\d+)\.\s+(.+)', stripped)
    if match:
        return ('numbered_item', int(match.group(1)), match.group(2))

    # Regular paragraph
    return ('paragraph', 0, stripped)


def apply_inline_formatting(paragraph, text):
    """Apply bold, italic, and other inline formatting to a paragraph."""
    # Clean HTML entities first
    text = clean_html_entities(text)

    # Split on bold/italic markers and apply
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def add_markdown_content(doc, markdown_text, is_front_matter=False):
    """Convert markdown text to Word document content."""
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    current_paragraph_lines = []

    def flush_paragraph():
        nonlocal current_paragraph_lines
        if current_paragraph_lines:
            text = ' '.join(current_paragraph_lines)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            apply_inline_formatting(p, text)
            current_paragraph_lines = []

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            # Parse table
            header = [cell.strip() for cell in table_rows[0].strip('|').split('|')]
            data_rows = []
            for row in table_rows[2:]:  # Skip header separator
                cells = [cell.strip() for cell in row.strip('|').split('|')]
                data_rows.append(cells)

            if header and data_rows:
                ncols = len(header)
                table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
                table.style = 'Table Grid'

                # Header row
                for j, h in enumerate(header):
                    if j < ncols:
                        cell = table.rows[0].cells[j]
                        cell.text = strip_markdown_formatting(h)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
                            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

                # Data rows
                for ri, row_data in enumerate(data_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < ncols:
                            cell = table.rows[ri + 1].cells[j]
                            cell.text = strip_markdown_formatting(cell_text)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(10)
                                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            table_rows = []
            in_table = False

    while i < len(lines):
        line = lines[i]
        ptype, level, text = parse_markdown_line(line)

        if ptype == 'table_row':
            flush_paragraph()
            if not in_table:
                in_table = True
            table_rows.append(line.strip())
            i += 1
            continue
        elif in_table:
            flush_table()

        if ptype == 'hr':
            flush_paragraph()
            # Skip horizontal rules - they're section separators in markdown

        elif ptype == 'heading':
            flush_paragraph()
            if level <= 5:
                heading = doc.add_heading(clean_html_entities(text), level=level)
                # Ensure TNR font
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.color.rgb = RGBColor(0, 0, 0)

        elif ptype == 'empty':
            flush_paragraph()

        elif ptype == 'list_item':
            flush_paragraph()
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, text)

        elif ptype == 'numbered_item':
            flush_paragraph()
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, text)

        elif ptype == 'paragraph':
            current_paragraph_lines.append(text)

        i += 1

    flush_paragraph()
    if in_table:
        flush_table()


def add_figure(doc, figure_path, caption, figure_num):
    """Add a figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if figure_path.exists():
        run = p.add_run()
        run.add_picture(str(figure_path), width=Inches(5.5))
    else:
        p.add_run(f'[Figure {figure_num}: {figure_path.name} - file not found]')

    # Caption
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(12)
    run = caption_p.add_run(f'Figure {figure_num}. ')
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = caption_p.add_run(caption)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def build_dissertation():
    """Build the complete dissertation Word document."""
    print("=" * 60)
    print("BUILDING DISSERTATION DOCUMENT")
    print("=" * 60)

    doc = create_document()

    # Chapter files in order
    chapter_files = [
        ('00_front_matter.md', True),
        ('01_introduction.md', False),
        ('02_literature_review.md', False),
        ('03_methodology.md', False),
        ('04_results.md', False),
        ('05_discussion.md', False),
        ('06_references.md', False),
        ('07_appendices.md', False),
    ]

    for filename, is_front in chapter_files:
        filepath = CHAPTERS_DIR / filename
        if filepath.exists():
            print(f"  Adding: {filename}")
            content = filepath.read_text(encoding='utf-8')

            # Add page break before each chapter (except front matter)
            if filename != '00_front_matter.md':
                add_page_break(doc)

            add_markdown_content(doc, content, is_front_matter=is_front)
        else:
            print(f"  MISSING: {filename}")

    # Insert figures in Results chapter area
    figures = [
        ('fig_4_1_procurement_design_distribution.png', 'Distribution of Awards by Procurement Design Category', '4.1'),
        ('fig_4_2_award_amount_distribution.png', 'Distribution of Award Amounts by Procurement Design', '4.2'),
        ('fig_4_3_cost_growth_boxplot.png', 'Cost Growth Distribution by Procurement Design', '4.3'),
        ('fig_4_4_single_bid_rate.png', 'Single-Bid Rate by Procurement Design Category', '4.4'),
        ('fig_4_5_offers_distribution.png', 'Distribution of Number of Offers Received', '4.5'),
        ('fig_4_6_psm_distribution.png', 'Propensity Score Distribution Before and After Matching', '4.6'),
        ('fig_4_7_balance_plot.png', 'Covariate Balance Before and After Matching', '4.7'),
        ('fig_4_8_marginal_effects.png', 'Marginal Effect of Quality-Evaluating Design by Award Size', '4.8'),
    ]

    # Add figures appendix
    add_page_break(doc)
    doc.add_heading('Figures', level=1)
    for fig_file, caption, num in figures:
        fig_path = FIGURES_DIR / fig_file
        add_figure(doc, fig_path, caption, num)

    # Save
    output_path = OUTPUT_DIR / 'Davidson_Dissertation_Draft.docx'
    doc.save(str(output_path))
    print(f"\nSaved: {output_path}")
    print(f"Size: {output_path.stat().st_size / 1024:.0f} KB")

    return output_path


if __name__ == '__main__':
    build_dissertation()
