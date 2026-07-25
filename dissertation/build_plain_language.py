#!/usr/bin/env python3
"""
Build plain-language Word documents from markdown files.
Produces three versions: Policy Brief, White Paper, Full Rewrite.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

OUTPUT_DIR = Path(__file__).parent / 'output'


def create_doc():
    """Create a clean Word document with professional formatting."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    for level in range(1, 6):
        sname = f'Heading {level}'
        if sname in doc.styles:
            hs = doc.styles[sname]
            hf = hs.font
            hf.name = 'Calibri'
            hf.color.rgb = RGBColor(0, 51, 102)
            if level == 1:
                hf.size = Pt(18)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(18)
                hs.paragraph_format.space_after = Pt(10)
            elif level == 2:
                hf.size = Pt(14)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(14)
                hs.paragraph_format.space_after = Pt(6)
            elif level == 3:
                hf.size = Pt(12)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(10)
                hs.paragraph_format.space_after = Pt(4)
            else:
                hf.size = Pt(11)
                hf.bold = True
                hs.paragraph_format.space_before = Pt(8)
                hs.paragraph_format.space_after = Pt(4)

    return doc


def clean_text(text):
    """Clean HTML entities and special characters."""
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&mdash;', '\u2014')
    text = text.replace('&ndash;', '\u2013')
    return text


def apply_formatting(paragraph, text):
    """Apply bold/italic inline formatting."""
    text = clean_text(text)
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            paragraph.add_run(part[1:-1])
        else:
            paragraph.add_run(part)


def strip_formatting(text):
    """Remove markdown formatting markers."""
    text = clean_text(text)
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    doc = create_doc()
    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_table = False
    table_rows = []
    para_lines = []
    in_blockquote = False
    bq_lines = []

    def flush_para():
        nonlocal para_lines
        if para_lines:
            text = ' '.join(para_lines)
            p = doc.add_paragraph()
            apply_formatting(p, text)
            para_lines = []

    def flush_bq():
        nonlocal bq_lines, in_blockquote
        if bq_lines:
            text = ' '.join(bq_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            # Light gray background via shading
            apply_formatting(p, text)
            for run in p.runs:
                run.italic = True
            bq_lines = []
            in_blockquote = False

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            header = [c.strip() for c in table_rows[0].strip('|').split('|')]
            data = []
            for row in table_rows[2:]:
                cells = [c.strip() for c in row.strip('|').split('|')]
                data.append(cells)

            if header and data:
                ncols = len(header)
                table = doc.add_table(rows=1 + len(data), cols=ncols)
                table.style = 'Light Shading Accent 1'

                for j, h in enumerate(header):
                    if j < ncols:
                        cell = table.rows[0].cells[j]
                        cell.text = strip_formatting(h)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.bold = True
                                r.font.size = Pt(10)

                for ri, rd in enumerate(data):
                    for j, ct in enumerate(rd):
                        if j < ncols:
                            cell = table.rows[ri + 1].cells[j]
                            cell.text = strip_formatting(ct)
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(10)

            table_rows = []
            in_table = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            flush_para()
            flush_bq()
            if in_table:
                flush_table()
            i += 1
            continue

        # Table row
        if stripped.startswith('|') and stripped.endswith('|'):
            flush_para()
            flush_bq()
            if not in_table:
                in_table = True
            table_rows.append(stripped)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Heading
        if stripped.startswith('#'):
            flush_para()
            flush_bq()
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            if level <= 5:
                h = doc.add_heading(clean_text(text), level=level)
                for run in h.runs:
                    run.font.name = 'Calibri'
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            flush_para()
            in_blockquote = True
            bq_text = stripped[1:].strip()
            if bq_text:
                bq_lines.append(bq_text)
            i += 1
            continue
        elif in_blockquote and stripped:
            bq_lines.append(stripped)
            i += 1
            continue
        elif in_blockquote:
            flush_bq()

        # Empty line
        if not stripped:
            flush_para()
            i += 1
            continue

        # List items
        if stripped.startswith('- ') or stripped.startswith('* '):
            flush_para()
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting(p, stripped[2:])
            i += 1
            continue

        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            flush_para()
            p = doc.add_paragraph(style='List Number')
            apply_formatting(p, m.group(2))
            i += 1
            continue

        # Regular paragraph
        para_lines.append(stripped)
        i += 1

    flush_para()
    flush_bq()
    if in_table:
        flush_table()

    doc.save(str(docx_path))
    size = docx_path.stat().st_size / 1024
    print(f"  Saved: {docx_path} ({size:.0f} KB)")


def main():
    files = [
        ('Policy_Brief.md', 'Davidson_Policy_Brief.docx'),
        ('White_Paper.md', 'Davidson_White_Paper.docx'),
        ('Full_Plain_Language.md', 'Davidson_Full_Plain_Language.docx'),
        ('Global_Procurement_White_Paper.md', 'Davidson_Global_Procurement_White_Paper.docx'),
        ('Global_Procurement_Full_Report.md', 'Davidson_Global_Procurement_Full_Report.docx'),
        ('Global_Procurement_Exhaustive_Treatise.md', 'Davidson_Global_Procurement_Exhaustive_Treatise.docx'),
    ]

    print("Building Word documents...")
    for md_name, docx_name in files:
        md_path = OUTPUT_DIR / md_name
        docx_path = OUTPUT_DIR / docx_name
        if md_path.exists():
            print(f"  Converting: {md_name}")
            md_to_docx(md_path, docx_path)
        else:
            print(f"  MISSING: {md_name}")

    print("Done.")


if __name__ == '__main__':
    main()
